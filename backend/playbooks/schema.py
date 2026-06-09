import uuid
from services.publisher import get_publisher
from django.utils import timezone
from core.rabbitmq import publish_event
import json
import base64
import posixpath
import re
import yaml
import requests
from graphene.types.generic import GenericScalar
from graphene_file_upload.scalars import Upload
from django.db.models import Count, Q, Prefetch
from playbooks.models import DetectionPlaybook
from platform_data.models import MitreAttackTechnique
import graphene
from graphene_django import DjangoObjectType
from .models import (
    DetectionPlaybook,
    Task,
    PlaybookGraph,
    PlaybookNode,
    PlaybookEdge,
    ActivityLog,
    PlaybookComment,
    CapabilityAbstraction,
)
from review.models import ReviewRequest as CanonReviewRequest, ReviewComment as CanonReviewComment
from tags.schema import TagType
from tags.models import TenantTag
from rules.models import DetectionRule, RuleRepository
from data_catalog.models import DataSource
from rules.schema import RuleType
from data_catalog.schema import DataSourceType
from identity.decorators import role_required, Roles
from identity.schema import UserType
from review.schema import ReviewRequestType, ReviewCommentType
from organizations.schema import OrganizationType
from platform_data.schema import D3fendTechniqueType
from platform_data.models import (
    MitreAttackTechnique,
    MitreIcsTechnique,
    MitreMobileTechnique,
    D3fendDefensiveTechnique,
)
import logging

export_logger = logging.getLogger(__name__)


def generate_copy_title(base_title: str, existing_titles) -> str:
    """Return a unique COPY title with incremental numeric suffixes."""
    suffix_base = f"{base_title} - COPY"
    if suffix_base not in existing_titles:
        return suffix_base
    counter = 1
    while True:
        candidate = f"{suffix_base}{counter}"
        if candidate not in existing_titles:
            return candidate
        counter += 1


def _notify_dac_automation_failure(graph, actor, error_message: str) -> None:
    try:
        from django.contrib.contenttypes.models import ContentType
        from notifications.models import Notification

        recipient = actor if getattr(actor, 'id', None) else getattr(graph, 'author', None)
        if recipient is None:
            return

        Notification.objects.create(
            recipient=recipient,
            actor=actor if getattr(actor, 'id', None) else None,
            organization=graph.organization,
            verb=f"DaC deployment automation failed: {error_message}",
            object_id=str(graph.id),
            content_type=ContentType.objects.get_for_model(PlaybookGraph),
        )
    except Exception:
        export_logger.exception(
            'Failed to create notification for DaC automation failure graph_id=%s',
            getattr(graph, 'id', None),
        )


def _queue_dac_deployment_automation(graph, actor) -> bool:
    from organizations.models import DacDeploymentConfig, OpenTideHefPublishJob, PlatformCredential
    from playbooks.utils.opentide_compiler import (
        _normalize_mdr_impacted_entities,
        compile_mdr_yaml,
        dump_opentide_yaml,
    )
    from rules.deployers import PLATFORM_DEPLOYER_MAP
    from rules.opentide_publish import (
        deploy_opentide_rule_to_platforms,
        upsert_opentide_rule_for_graph,
    )

    config = DacDeploymentConfig.objects.select_related(
        'target_repository',
        'publish_profile',
    ).filter(
        organization=graph.organization,
    ).first()
    if not config or config.mode == DacDeploymentConfig.Mode.NONE:
        return False

    requested_platforms = []
    if config.mode in (
        DacDeploymentConfig.Mode.GIT_PUSH_AND_DEPLOY,
        DacDeploymentConfig.Mode.DEPLOY_ONLY,
    ):
        requested_platforms = [str(p).strip().lower() for p in (config.target_platforms or []) if str(p).strip()]
        valid_platforms = set(PLATFORM_DEPLOYER_MAP.keys())
        unsupported = [p for p in requested_platforms if p not in valid_platforms]
        if unsupported:
            raise RuntimeError(f"Unsupported DaC target platforms: {', '.join(unsupported)}")
        requested_platforms = list(dict.fromkeys(requested_platforms))
        if not requested_platforms:
            raise RuntimeError(f'No target platforms configured for {config.mode} mode')
        configured_creds = set(
            PlatformCredential.objects.filter(
                organization=graph.organization,
                platform__in=requested_platforms,
                enabled=True,
            ).values_list('platform', flat=True)
        )
        missing = [platform for platform in requested_platforms if platform not in configured_creds]
        if missing:
            raise RuntimeError(
                f"Missing enabled platform credentials for: {', '.join(missing)}"
            )

    if config.mode == DacDeploymentConfig.Mode.DEPLOY_ONLY:
        source = 'DAC_DEPLOY_ONLY'
        source_details = f'{source}:{graph.version}:{graph.minor_version}'
        existing_trigger = ActivityLog.objects.filter(
            playbook=graph,
            action=source,
            details=source_details,
        ).exists()
        if existing_trigger:
            export_logger.info(
                'Skipping DaC deploy-only automation for graph %s: already executed for version %s.%s',
                graph.id,
                graph.version,
                graph.minor_version,
            )
            return False

        actor_for_rule = actor if getattr(actor, 'id', None) else getattr(graph, 'author', None)
        if actor_for_rule is None:
            raise RuntimeError('No actor available for DaC deploy-only automation')

        ActivityLog.objects.create(
            playbook=graph,
            user=actor_for_rule if getattr(actor_for_rule, 'id', None) else None,
            action=source,
            details=source_details,
        )

        mdr_data = compile_mdr_yaml(graph)
        _normalize_mdr_impacted_entities(mdr_data)
        raw_yaml = dump_opentide_yaml(mdr_data)
        rule = upsert_opentide_rule_for_graph(
            graph,
            actor_for_rule,
            raw_yaml,
            repository=None,
        )
        deployment_results, overall_success, _message = deploy_opentide_rule_to_platforms(
            rule,
            graph.organization,
            requested_platforms,
        )
        if not overall_success:
            failed_platforms = [
                result['platform']
                for result in deployment_results
                if not result.get('success')
            ]
            raise RuntimeError(
                'Deployment failed for platform(s): ' + ', '.join(failed_platforms or requested_platforms)
            )

        export_logger.info(
            'Completed DaC deploy-only automation for graph %s platforms=%s',
            graph.id,
            requested_platforms,
        )
        return True

    existing_trigger = OpenTideHefPublishJob.objects.filter(
        playbook=graph,
        source='DAC_AUTOMATION',
        source_graph_version=graph.version,
        source_graph_minor_version=graph.minor_version,
    ).exists()
    if existing_trigger:
        export_logger.info(
            'Skipping DaC automation for graph %s: already queued for version %s.%s',
            graph.id,
            graph.version,
            graph.minor_version,
        )
        return False

    repository = config.target_repository
    if repository is None:
        raise RuntimeError('No target repository configured for DaC deployment automation')
    if not repository.git_url:
        raise RuntimeError('Configured DaC target repository has no git URL')
    if not repository.token:
        raise RuntimeError('Configured DaC target repository has no access token')

    job = OpenTideHefPublishJob.objects.create(
        playbook=graph,
        user=actor if getattr(actor, 'id', None) else getattr(graph, 'author', None),
        organization=graph.organization,
        profile=config.publish_profile,
        repository=repository,
        status='QUEUED',
        commit_message=f'Auto-publish on DEPLOYED: {graph.title}',
        branch=(config.target_branch or 'main'),
        target_folder=config.target_folder or '',
        push_opentide_bundle=True,
        push_platform_rules=True,
        requested_platforms=requested_platforms,
        source='DAC_AUTOMATION',
        source_graph_version=graph.version,
        source_graph_minor_version=graph.minor_version,
    )

    published = publish_event(
        'opentide.hef.publish.queued',
        {
            'task_id': str(job.id),
            'playbook_id': str(graph.id),
            'organization_id': str(graph.organization.id),
            'user_id': str(actor.id) if getattr(actor, 'id', None) else '',
        },
    )
    if not published:
        job.status = 'FAILED'
        job.error_message = 'Failed to queue DaC automation publish job: RabbitMQ unavailable.'
        job.save(update_fields=['status', 'error_message'])
        raise RuntimeError('Failed to queue DaC automation publish job: RabbitMQ unavailable')

    export_logger.info(
        'Queued DaC automation HEF publish job %s for graph %s mode=%s platforms=%s',
        job.id,
        graph.id,
        config.mode,
        requested_platforms,
    )
    return True


# --- Platform Data GraphQL Types ---
class MitreAttackTechniqueType(DjangoObjectType):
    # Expose the UUID id as a proper UUID type instead of relay ID
    id = graphene.UUID(source='pk')
    
    class Meta:
        model = MitreAttackTechnique
        fields = "__all__"


class MitreIcsTechniqueType(DjangoObjectType):
    class Meta:
        model = MitreIcsTechnique
        fields = "__all__"


class MitreMobileTechniqueType(DjangoObjectType):
    class Meta:
        model = MitreMobileTechnique
        fields = "__all__"


class CapabilityAbstractionType(DjangoObjectType):
    technique = graphene.Field(MitreAttackTechniqueType)
    is_editable = graphene.Boolean()
    is_shared_baseline = graphene.Boolean()
    organization_name = graphene.String()

    class Meta:
        model = CapabilityAbstraction
        fields = (
            "id",
            "technique",
            "organization",
            "created_by",
            "updated_by",
            "abstraction_layer",
            "component_artifact",
            "adversary_purpose",
            "common_evasions",
            "expected_observables",
            "applicable_telemetry",
            "detection_value",
            "robustness_level",
            "source_kind",
            "review_status",
            "is_baseline",
            "version",
            "created_at",
            "updated_at",
        )

    def resolve_is_editable(self, info):
        user = info.context.user
        return bool(
            user
            and not user.is_anonymous
            and self.organization_id
            and self.organization_id == getattr(user.organization, 'id', None)
        )

    def resolve_is_shared_baseline(self, info):
        return self.organization_id is None

    def resolve_organization_name(self, info):
        return self.organization.name if self.organization else "Shared Baseline"

    def resolve_technique(self, info):
        return self.technique

# --- Object Types ---
class ActivityLogType(DjangoObjectType):
    class Meta:
        model = ActivityLog
        fields = ("id", "user", "action", "details", "timestamp")

class PlaybookCommentType(DjangoObjectType):
    class Meta:
        model = PlaybookComment
        fields = "__all__"

class TaskType(DjangoObjectType):
    class Meta:
        model = Task
        fields = "__all__"

## Use shared Review types from review.schema to avoid duplicate type names

class PlaybookType(DjangoObjectType):
    # Existing relations as lists
    tags = graphene.List(TagType)
    detection_rules = graphene.List(RuleType)
    required_data_sources = graphene.List(DataSourceType)
    tasks = graphene.List(TaskType)

    # --- SECTION 1: METADATA ---
    # CamelCase aliases for UI consistency (source maps to model field names)
    analyticId = graphene.String(source='analytic_id', description="Unique ID for the analytic (e.g., DE-2025-001)")
    version = graphene.String(description="Semantic version of the analytic (e.g., 1.0)")

    # --- SECTION 2: DETECTION OVERVIEW ---
    hypothesis = graphene.String(description="The testable question this hunt is designed to answer")

    # --- SECTION 3: ANALYTIC DETAILS ---
    robustnessLevel = graphene.Int(source='robustness_level', description="Level 1-5 of how easy this detection is to evade")
    dataSourceRobustness = graphene.String(source='data_source_robustness', description="Where does the data come from? (e.g., K for Kernel-Mode)")

    # --- SECTION 4: LOGIC & IMPLEMENTATION ---
    falsePositiveRate = graphene.Int(source='false_positive_rate', description="Estimated FP rate on a 0-100% scale")
    knownFalsePositives = graphene.String(source='known_false_positives', description="Known benign triggers to reduce noise")
    exclusionStrategy = graphene.String(source='exclusion_strategy', description="Surgically precise exclusions (e.g., trusted signers)")

    # --- SECTION 5: VALIDATION & RESPONSE ---
    testingProcedures = graphene.String(source='testing_procedures', description="Test cases and synonym tools to validate the detection")
    triageGuidance = graphene.String(source='triage_guidance', description="Step-by-step instructions for T1 analysts")

    # --- SECTION 6: AUTOMATION (SOAR) ---
    soarEnrichment = graphene.String(source='soar_enrichment', description="Automated enrichment steps")
    soarTriage = graphene.String(source='soar_triage', description="Automated triage logic")
    soarContainment = graphene.String(source='soar_containment', description="Automated containment steps")

    # --- CAPABILITY ABSTRACTION ---
    operationalPath = graphene.String(source='operational_path', description="High-level operational path / phases")
    functionCallGraphs = graphene.String(source='function_call_graphs', description="Relevant function/API call graph abstractions")
    executionModalities = graphene.String(source='execution_modalities', description="Execution modalities or behavioral variants")
    # --- HUNT TECHNICAL DETAILS ---
    technicalDetails = graphene.String(source='technical_details', description="Deep technical implementation notes (for HUNT playbooks)")

    # --- SECTION 7: FRAMEWORK MAPPINGS ---
    mitreAttackMappings = graphene.List(MitreAttackTechniqueType)
    mitreIcsMappings = graphene.List(MitreIcsTechniqueType)
    mitreMobileMappings = graphene.List(MitreMobileTechniqueType)

    # --- SHARING / OWNERSHIP METADATA ---
    owner_organization_name = graphene.String(description="The name of the organization that owns this playbook.")
    is_read_only = graphene.Boolean(description="True if the playbook is shared and not owned by the user's org.")

    # Attached abstraction capability graphs
    graphs = graphene.List(lambda: PlaybookGraphType)

    class Meta:
        model = DetectionPlaybook
        # Explicitly include only scalar/simple fields. Relations (tags, M2M) are
        # exposed via explicit GraphQL fields and resolvers above.
        fields = (
            "id",
            "title",
            "description",
            "status",
            "playbook_type",
            "author",
            "organization",
            "created_at",
            "updated_at",
            # detection-tmpl fields
            "analytic_id",
            "version",
            "hypothesis",
            "robustness_level",
            "data_source_robustness",
            "false_positive_rate",
            "known_false_positives",
            "exclusion_strategy",
            "testing_procedures",
            "triage_guidance",
            "soar_enrichment",
            "soar_triage",
            "soar_containment",
            "operational_path",
            "function_call_graphs",
            "execution_modalities",
            "technical_details",
            "is_shared",
        )

    def resolve_tags(self, info):
        return self.tags.all()

    def resolve_detection_rules(self, info):
        return self.detection_rules.all()

    def resolve_required_data_sources(self, info):
        return self.required_data_sources.all()

    def resolve_tasks(self, info):
        # Tasks have been moved to PlaybookGraph. 
        # For legacy DetectionPlaybook, we return an empty list to avoid errors.
        return []

    # --- NEW RESOLVERS FOR M2M LINKS ---
    def resolve_mitreAttackMappings(self, info):
        return self.mitre_attack_mappings.all()

    def resolve_mitreIcsMappings(self, info):
        return self.mitre_ics_mappings.all()

    def resolve_mitreMobileMappings(self, info):
        return self.mitre_mobile_mappings.all()

    # --- SHARING / OWNERSHIP RESOLVERS ---
    def resolve_owner_organization_name(self, info):
        # self is a DetectionPlaybook instance
        if self.organization:
            return self.organization.name
        return None

    def resolve_is_read_only(self, info):
        user = info.context.user
        if user.is_anonymous:
            # Default to read-only if we don't know who the user is
            return True
        # If the playbook's org is not my org, it's read-only
        return self.organization != user.organization

    def resolve_graphs(self, info):
        return self.graphs.all()

# --- NEW V2 TYPES (GRAPH-BASED) ---

class PlaybookNodeType(DjangoObjectType):
    ui_metadata = graphene.JSONString()
    template_data = graphene.JSONString()
    mitre_attack_mappings = graphene.List(MitreAttackTechniqueType)

    class Meta:
        model = PlaybookNode
        fields = (
            "id",
            "graph",
            "layer_name",
            "position_x",
            "position_y",
            "ui_metadata",
            "color",
            "template_data",
        )

    def resolve_ui_metadata(self, info):
        return self.ui_metadata
    
    def resolve_template_data(self, info):
        return self.template_data
    
    def resolve_mitre_attack_mappings(self, info):
        return self.mitre_attack_mappings.all()

class PlaybookEdgeType(DjangoObjectType):
    # React Flow needs 'source' and 'target' IDs, not full objects
    source = graphene.String()
    target = graphene.String()

    class Meta:
        model = PlaybookEdge
        fields = ("id", "source", "target")

    def resolve_source(self, info):
        return self.source_node_id

    def resolve_target(self, info):
        return self.target_node_id

class PlaybookGraphType(DjangoObjectType):
    nodes = graphene.List(PlaybookNodeType)
    edges = graphene.List(PlaybookEdgeType)

    owner_organization_name = graphene.String()
    organization = graphene.Field(OrganizationType)
    is_read_only = graphene.Boolean()

    # Notes, attached playbooks, and PNG snapshot URL
    notes = graphene.String()
    playbooks = graphene.List(PlaybookType)
    pngSnapshotUrl = graphene.String()
    graph_image_url = graphene.String()
    
    # Add resolvers for the new relationships
    tags = graphene.List(graphene.String) # Simple string list for tags
    tasks = graphene.List(TaskType)
    activities = graphene.List(ActivityLogType)
    comments = graphene.List(PlaybookCommentType)
    active_review = graphene.Field(ReviewRequestType)

    # --- NEW STRATEGY FIELDS ---
    mitre_technique = graphene.Field(MitreAttackTechniqueType)
    selected_strategy = graphene.JSONString()
    detection_rule = graphene.String()
    
    # --- D3FEND DEFENSIVE TECHNIQUES ---
    d3fend_techniques = graphene.List(D3fendTechniqueType)
    selected_capability_abstractions = graphene.List(CapabilityAbstractionType)
    detection_focus_layer = graphene.String()
    
    # --- NEW CONTEXT FIELDS ---
    goal = graphene.String()
    technical_context = graphene.String()
    blind_spots = graphene.String()
    triage_guidance = graphene.String()
    false_positives = graphene.String()
    response_playbook = graphene.String()

    # --- DEPLOYMENT METADATA ---
    target_file_path = graphene.String()
    git_status = graphene.String()
    last_commit_hash = graphene.String()

    # --- METADATA (DCG420) ---
    custom_id = graphene.String()
    version = graphene.Int()
    minor_version = graphene.Int()
    robustness_level = graphene.Int()
    data_source_robustness = graphene.String()
    
    # Maieutic Engine fields
    data_source_maturity = graphene.String()
    conversation_history = graphene.JSONString()
    
    # SOAR Fields
    alert_trigger = graphene.String()
    default_severity = graphene.String()
    enrichment_steps = graphene.JSONString()
    containment_steps = graphene.JSONString()
    notification_steps = graphene.JSONString()
    downstream_correlation_requirements = graphene.JSONString()

    # OpenTide Fields
    opentide_yaml = graphene.JSONString()
    configured_platforms = graphene.List(graphene.String)

    # OpenTide v2.1 Fields
    tlp_classification = graphene.String()
    public_references = graphene.JSONString()
    internal_references = graphene.JSONString()
    threat_actors = graphene.JSONString()
    threat_surface = graphene.JSONString()

    class Meta:
        model = PlaybookGraph
        fields = (
            "id", "title", "status", "is_shared", "author", "organization", "updated_at", "created_at",
            "nodes", "edges", "notes", "playbooks", "png_snapshot",
            "mitre_technique", "selected_strategy", "detection_rule",
             "goal", "technical_context", "blind_spots", "triage_guidance",
             "false_positives", "response_playbook",
             "test_scenario", "test_expected_output",
             "target_file_path", "git_status", "last_commit_hash",
             "custom_id", "version", "minor_version", "robustness_level", "data_source_robustness",
             "data_source_maturity", "conversation_history",
             "selected_capability_abstractions", "detection_focus_layer",
             "alert_trigger", "default_severity", "enrichment_steps", "containment_steps", "notification_steps",
             "downstream_correlation_requirements",
             "opentide_yaml", "configured_platforms",
             "tlp_classification", "public_references", "internal_references", "threat_actors", "threat_surface",
         )

    def resolve_tags(self, info):
        return self.tags.names() # Returns list of tag names

    def resolve_tasks(self, info):
        return self.tasks.all()

    def resolve_activities(self, info):
        return self.activities.all().order_by('-timestamp')

    def resolve_comments(self, info):
        return self.comments.all()

    def resolve_active_review(self, info):
        # Resolve the latest OPEN canonical review request for any attached playbook
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            from review.models import ReviewRequest as CanonReviewRequest
        except Exception:
            logger.warning(f"[resolve_active_review] Failed to import CanonReviewRequest")
            return None

        playbook_ids = list(self.playbooks.values_list('id', flat=True))
        logger.info(f"[resolve_active_review] Graph {self.id}: playbook_ids={playbook_ids}")
        
        if not playbook_ids:
            logger.info(f"[resolve_active_review] No playbooks attached to graph {self.id}")
            return None
            
        req = CanonReviewRequest.objects.filter(
            playbook_id__in=playbook_ids,
            status=CanonReviewRequest.ReviewStatus.OPEN
        ).order_by('created_at').last()
        
        if req:
            comments_count = req.comments.count()
            logger.info(f"[resolve_active_review] Found active review {req.id} with {comments_count} comments")
        else:
            logger.info(f"[resolve_active_review] No OPEN review request found for playbook_ids={playbook_ids}")
        
        return req

    def resolve_nodes(self, info):
        return self.nodes.all()
    
    def resolve_selected_strategy(self, info):
        return self.selected_strategy

    def resolve_edges(self, info):
        return self.edges.all()

    def resolve_owner_organization_name(self, info):
        return self.organization.name

    def resolve_is_read_only(self, info):
        user = info.context.user
        return self.organization != user.organization

    def resolve_notes(self, info):
        return self.notes

    def resolve_playbooks(self, info):
        return self.playbooks.all()

    def resolve_d3fend_techniques(self, info):
        return self.d3fend_techniques.all()

    def resolve_selected_capability_abstractions(self, info):
        return self.selected_capability_abstractions.select_related('organization', 'technique')

    def resolve_pngSnapshotUrl(self, info):
        if self.png_snapshot:
            try:
                url = info.context.build_absolute_uri(self.png_snapshot.url)
                # Force HTTPS if the request came over HTTPS
                if info.context.is_secure():
                    url = url.replace('http://', 'https://')
                return url
            except Exception:
                return self.png_snapshot.url
        return None

    def resolve_graph_image_url(self, info):
        if self.png_snapshot:
            # Return the full URL to the image
            try:
                url = info.context.build_absolute_uri(self.png_snapshot.url)
                # Force HTTPS if the request came over HTTPS
                if info.context.is_secure():
                    url = url.replace('http://', 'https://')
                return url
            except Exception:
                return self.png_snapshot.url
        return None

# --- Query Definitions ---

class Query(graphene.ObjectType):
    playbook_meta = graphene.Field(graphene.JSONString, description="Consolidated choices and technique lists for playbooks")
    all_playbooks = graphene.List(PlaybookType, description="Retrieves all playbooks for the user's organization.")
    playbook = graphene.Field(PlaybookType, id=graphene.UUID(required=True), description="Retrieves a single playbook by its UUID, ensuring it belongs to the user's organization.")
    
    attack_navigator_layer = graphene.JSONString(
        description="Generates a JSON layer for the ATT&CK Navigator (v2 Graph Model).",
        include_d3fend=graphene.Boolean(default_value=False)
    )

    # We will create a new v2 query.
    attack_navigator_layer_v2 = graphene.JSONString(
        description="Generates a JSON layer for the ATT&CK Navigator (v2 Graph Model).",
        include_d3fend=graphene.Boolean(default_value=False)
    )
    # --- Platform data lists for creation forms ---
    all_attack_techniques = graphene.List(
        MitreAttackTechniqueType,
        description="All MITRE ATT&CK techniques (global)",
        search=graphene.String(),
        limit=graphene.Int(default_value=50),
        offset=graphene.Int(default_value=0),
    )
    all_ics_techniques = graphene.List(
        MitreIcsTechniqueType,
        description="All MITRE ICS techniques (global)",
        search=graphene.String(),
        limit=graphene.Int(default_value=50),
        offset=graphene.Int(default_value=0),
    )
    all_mobile_techniques = graphene.List(
        MitreMobileTechniqueType,
        description="All MITRE Mobile techniques (global)",
        search=graphene.String(),
        limit=graphene.Int(default_value=50),
        offset=graphene.Int(default_value=0),
    )

    # --- NEW V2 QUERIES ---
    playbook_graph = graphene.Field(
        PlaybookGraphType, 
        id=graphene.UUID(required=True),
        description="Retrieves a single playbook graph (v2) by its ID."
    )

    all_playbook_graphs = graphene.List(
        PlaybookGraphType,
        description="Retrieves all v2 playbook graphs for the user's organization."
    )
    capability_abstractions = graphene.List(
        CapabilityAbstractionType,
        technique_id=graphene.String(),
        include_baseline=graphene.Boolean(default_value=True),
        description="Capability abstraction entries for an ATT&CK technique, including shared baseline and organization-specific content.",
    )

    def resolve_all_playbooks(self, info):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # 1. Get playbooks my organization owns
        my_org_playbooks = Q(organization=user.organization)

        # 2. Get playbooks shared with me (if my org is in an Entity)
        shared_playbooks = Q()  # Start with an empty (blank) query

        if getattr(user, "organization", None) and getattr(user.organization, "entity", None):
            # This Q object finds playbooks that are:
            # - In the same entity as me
            # - Have the 'is_shared' flag set to True
            # - AND (~) do not belong to my own organization
            shared_playbooks = Q(
                organization__entity=user.organization.entity,
                is_shared=True,
            ) & ~Q(organization=user.organization)

        # 3. Combine the queries with an OR operator (|)
        # Use .select_related('organization') to optimize the owner_organization_name resolver
        return DetectionPlaybook.objects.filter(
            my_org_playbooks | shared_playbooks
        ).distinct().select_related("organization")

    def resolve_playbook(self, info, id):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")
        try:
            playbook = DetectionPlaybook.objects.get(pk=id)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found")
        if playbook.organization != user.organization:
            raise Exception("You do not have permission to view this playbook")
        return playbook

    def resolve_attack_navigator_layer(self, info, include_d3fend=False):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # --- 1. Get DEPLOYED Graphs ---
        # We filter for Graphs that are marked DEPLOYED
        deployed_graphs = PlaybookGraph.objects.filter(
            organization=user.organization,
            status=DetectionPlaybook.PlaybookStatus.DEPLOYED
        ).prefetch_related('nodes__mitre_attack_mappings') # Optimize: Prefetch nodes and their techniques

        # --- 2. Aggregate Techniques from Nodes ---
        # We collect technique IDs from all nodes in all deployed graphs
        technique_counts = {} # { 'T1003.001': { count: 1, d3fend: set() } }

        for graph in deployed_graphs:
            for node in graph.nodes.all():
                for tech in node.mitre_attack_mappings.all():
                    tid = tech.technique_id

                    if tid not in technique_counts:
                        technique_counts[tid] = { 'count': 0, 'd3fend': set() }

                    technique_counts[tid]['count'] += 1

                    # Note: D3FEND logic would go here if we kept it. 
                    # Since we removed D3FEND models on Day 179, we skip that part.

        # --- 3. Format JSON for Navigator ---
        COVERAGE_COLOR = "#a1d99b"
        techniques_list = []
        max_score = 0

        for tid, data in technique_counts.items():
            count = data['count']
            if count > max_score: max_score = count

            techniques_list.append({
                "techniqueID": tid,
                "color": COVERAGE_COLOR,
                "comment": f"Covered by {count} detection node(s)",
                "enabled": True,
                "score": count,
            })

        if max_score == 0: max_score = 1

        return {
            "name": f"{user.organization.name} Coverage (v2)",
            "versions": { "attack": "v18", "navigator": "5.0", "layer": "4.5" },
            "domain": "enterprise-attack",
            "description": "Live coverage from Playbook Graph Nodes",
            "techniques": techniques_list,
        }

    def resolve_attack_navigator_layer_v2(self, info, include_d3fend=False):
        return self.resolve_attack_navigator_layer(info, include_d3fend=include_d3fend)

    # --- Platform data resolvers ---
    def resolve_all_attack_techniques(self, info, search=None, limit=50, offset=0):
        # Allow higher limits for autocomplete (up to 1000 for all MITRE ATT&CK techniques)
        limit = max(1, min(limit or 50, 1000))
        qs = MitreAttackTechnique.objects.all().order_by('technique_id')
        if search:
            qs = qs.filter(Q(technique_id__icontains=search) | Q(name__icontains=search))
        return qs[offset: offset + limit]

    def resolve_all_ics_techniques(self, info, search=None, limit=50, offset=0):
        limit = max(1, min(limit or 50, 200))
        qs = MitreIcsTechnique.objects.all()
        if search:
            qs = qs.filter(Q(technique_id__icontains=search) | Q(name__icontains=search))
        return qs[offset: offset + limit]

    def resolve_all_mobile_techniques(self, info, search=None, limit=50, offset=0):
        limit = max(1, min(limit or 50, 200))
        qs = MitreMobileTechnique.objects.all()
        if search:
            qs = qs.filter(Q(technique_id__icontains=search) | Q(name__icontains=search))
        return qs[offset: offset + limit]

    # --- NEW V2 RESOLVERS ---
    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.VIEWER, Roles.REVIEWER])
    def resolve_playbook_graph(self, info, id):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            # Security: Check for ownership OR sharing
            graph = PlaybookGraph.objects.select_related(
                'organization'
            ).prefetch_related(
                'nodes', 'edges', 'selected_capability_abstractions__organization', 'selected_capability_abstractions__technique'
            ).get(pk=id)

            is_owned = graph.organization == user.organization
            # If multi-entity sharing is not implemented, honor global sharing flag
            is_shared = graph.is_shared
            # Allow connector_svc to see all graphs for notifications/sync
            is_connector = getattr(user, 'username', '') == 'connector_svc'

            if is_owned or is_shared or is_connector:
                return graph
            else:
                raise PermissionError("You do not have permission to view this graph.")

        except PlaybookGraph.DoesNotExist:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"[resolve_playbook_graph] PlaybookGraph not found: id={id}")
            return None

    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.VIEWER, Roles.REVIEWER])
    def resolve_all_playbook_graphs(self, info):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        my_org_graphs = Q(organization=user.organization)

        # Include any globally shared graphs from other organizations
        shared_graphs = Q(is_shared=True) & ~Q(organization=user.organization)

        return PlaybookGraph.objects.filter(
            my_org_graphs | shared_graphs
        ).distinct().select_related('organization')

    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.VIEWER, Roles.REVIEWER])
    def resolve_capability_abstractions(self, info, technique_id=None, include_baseline=True):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        qs = CapabilityAbstraction.objects.select_related(
            'technique', 'organization', 'created_by', 'updated_by'
        )
        if technique_id:
            qs = qs.filter(technique__technique_id=technique_id)
        if include_baseline:
            qs = qs.filter(Q(organization=user.organization) | Q(organization__isnull=True))
        else:
            qs = qs.filter(organization=user.organization)
        return qs.order_by('technique__technique_id', 'abstraction_layer', 'component_artifact')

    def resolve_playbook_meta(self, info):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")
        # Choices from model enums
        robustness = [
            {"value": choice[0], "label": choice[1]} for choice in DetectionPlaybook.RobustnessLevel.choices
        ]
        event_src = [
            {"value": choice[0], "label": choice[1]} for choice in DetectionPlaybook.EventRobustness.choices
        ]
        # Status & type – assume enums exist or derive from constants
        statuses = [
            {"value": choice[0], "label": choice[1]} for choice in DetectionPlaybook.PlaybookStatus.choices
        ] if hasattr(DetectionPlaybook, 'PlaybookStatus') else []
        types = [
            {"value": 'DETECTION', "label": 'Detection'},
            {"value": 'HUNT', "label": 'Hunt'},
        ]
        # Technique lists (limited fields)
        attack = [
            {"id": str(t.id), "techniqueId": t.technique_id, "name": t.name} for t in MitreAttackTechnique.objects.only('id', 'technique_id', 'name').order_by('technique_id')[:500]
        ]
        ics = [
            {"id": str(t.id), "techniqueId": t.technique_id, "name": t.name} for t in MitreIcsTechnique.objects.only('id', 'technique_id', 'name').order_by('technique_id')[:500]
        ]
        mobile = [
            {"id": str(t.id), "techniqueId": t.technique_id, "name": t.name} for t in MitreMobileTechnique.objects.only('id', 'technique_id', 'name').order_by('technique_id')[:500]
        ]
        return json.dumps({
            "robustnessLevels": robustness,
            "eventRobustness": event_src,
            "statuses": statuses,
            "playbookTypes": types,
            "attackTechniques": attack,
            "icsTechniques": ics,
            "mobileTechniques": mobile,
        })

# --- Mutation Definitions (Must be defined BEFORE the root Mutation class) ---

class CreatePlaybook(graphene.Mutation):
    class Arguments:
        title = graphene.String(required=True)
        description = graphene.String()
        playbook_type = graphene.String(default_value='HUNT')

        # --- SECTION 1: METADATA ---
        analytic_id = graphene.String(description="Unique ID for the analytic")
        version = graphene.String(description="Semantic version (e.g., 1.0)")

        # --- SECTION 2: DETECTION OVERVIEW ---
        hypothesis = graphene.String(description="The testable question for a HUNT playbook")

        # --- SECTION 3: ANALYTIC DETAILS ---
        robustness_level = graphene.Int(description="Level 1-5 of how easy this detection is to evade")
        data_source_robustness = graphene.String(description="Where does the data come from? (e.g., K for Kernel-Mode)")

        # --- SECTION 4: LOGIC & IMPLEMENTATION ---
        false_positive_rate = graphene.Int(description="Estimated FP rate 0-100%")
        known_false_positives = graphene.String(description="Known benign triggers for a DETECTION playbook")
        exclusion_strategy = graphene.String(description="Surgically precise exclusions")

        # --- SECTION 5: VALIDATION & RESPONSE ---
        testing_procedures = graphene.String(description="Test cases and synonym tools")
        triage_guidance = graphene.String(description="Step-by-step instructions for T1 analysts")

        # --- SECTION 6: AUTOMATION (SOAR) ---
        soar_enrichment = graphene.String(description="Automated enrichment steps")
        soar_triage = graphene.String(description="Automated triage logic")
        soar_containment = graphene.String(description="Automated containment steps")
        operational_path = graphene.String(description="Operational path phases")
        function_call_graphs = graphene.String(description="Function call graph abstractions")
        execution_modalities = graphene.String(description="Execution modalities / behavioral variants")
        technical_details = graphene.String(description="Deep technical implementation notes (HUNT)")

        # --- SECTION 7: FRAMEWORK MAPPINGS ---
        mitre_attack_mappings = graphene.List(graphene.ID)
        mitre_d3fend_mappings = graphene.List(graphene.ID)
        mitre_engage_mappings = graphene.List(graphene.ID)
        # Extended
        mitre_ics_mappings = graphene.List(graphene.ID)
        mitre_mobile_mappings = graphene.List(graphene.ID)
    playbook = graphene.Field(PlaybookType)
    class Meta:
        description = "Creates a new detection playbook for the user's organization."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        title = kwargs.get('title')
        description = kwargs.get('description')
        playbook_type = kwargs.get('playbook_type', 'HUNT')

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        playbook = DetectionPlaybook(
            title=title,
            description=description,
            playbook_type=playbook_type.upper(),
            author=user,
            organization=user.organization,
        )

        # --- Handle M2M fields ---
        attack_ids = kwargs.pop('mitre_attack_mappings', None)
        d3fend_ids = kwargs.pop('mitre_d3fend_mappings', None)
        engage_ids = kwargs.pop('mitre_engage_mappings', None)
        ics_ids = kwargs.pop('mitre_ics_mappings', None)
        mobile_ids = kwargs.pop('mitre_mobile_mappings', None)

        # Handle simple fields
        for field, value in kwargs.items():
            setattr(playbook, field, value)

        playbook.save()  # Must save before adding M2M relationships

        # --- Set M2M relationships ---
        if attack_ids is not None:
            playbook.mitre_attack_mappings.set(attack_ids)
        if d3fend_ids is not None:
            playbook.mitre_d3fend_mappings.set(d3fend_ids)
        if engage_ids is not None:
            playbook.mitre_engage_mappings.set(engage_ids)
        if ics_ids is not None:
            playbook.mitre_ics_mappings.set(ics_ids)
        if mobile_ids is not None:
            playbook.mitre_mobile_mappings.set(mobile_ids)

        return CreatePlaybook(playbook=playbook)

class UpdatePlaybookLinks(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)
        detection_rule_ids = graphene.List(graphene.ID)
        data_source_ids = graphene.List(graphene.ID)
        # --- Frameworks ---
        mitre_attack_ids = graphene.List(graphene.ID)
        mitre_d3fend_ids = graphene.List(graphene.ID)
        mitre_engage_ids = graphene.List(graphene.ID)
        mitre_ics_ids = graphene.List(graphene.ID)
        mitre_mobile_ids = graphene.List(graphene.ID)

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Links a playbook to its related content (rules, data sources, and frameworks)."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        playbook_id = kwargs.get('playbook_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the playbook and verify it belongs to the user's org
        try:
            playbook = DetectionPlaybook.objects.get(pk=playbook_id, organization=user.organization)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission")

        # Handle linking detection rules if the argument was provided
        if 'detection_rule_ids' in kwargs:
            playbook.detection_rules.set(kwargs['detection_rule_ids'])

        # Handle linking data sources if the argument was provided
        if 'data_source_ids' in kwargs:
            playbook.required_data_sources.set(kwargs['data_source_ids'])

        # Framework M2M setters
        if 'mitre_attack_ids' in kwargs:
            playbook.mitre_attack_mappings.set(kwargs['mitre_attack_ids'])
        if 'mitre_d3fend_ids' in kwargs:
            playbook.mitre_d3fend_mappings.set(kwargs['mitre_d3fend_ids'])
        if 'mitre_engage_ids' in kwargs:
            playbook.mitre_engage_mappings.set(kwargs['mitre_engage_ids'])
        if 'mitre_ics_ids' in kwargs:
            playbook.mitre_ics_mappings.set(kwargs['mitre_ics_ids'])
        if 'mitre_mobile_ids' in kwargs:
            playbook.mitre_mobile_mappings.set(kwargs['mitre_mobile_ids'])

        playbook.save()
        return UpdatePlaybookLinks(playbook=playbook)

class UpdatePlaybookTags(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)
        tag_names = graphene.List(graphene.String, required=True)

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Sets the tags for a specific playbook. This will overwrite any existing tags. New tags will be created for the organization if they do not exist."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, playbook_id, tag_names, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        try:
            # First, get the playbook instance
            playbook = DetectionPlaybook.objects.get(pk=playbook_id)
        except DetectionPlaybook.DoesNotExist:
            # Do not reveal whether the object exists - return a generic permission/exists error
            raise Exception("Playbook not found or you do not have permission")

        # CRITICAL SECURITY CHECK: Verify the playbook belongs to the user's organization
        if playbook.organization != user.organization:
            # Do not reveal whether the object exists - return the same generic error
            raise Exception("Playbook not found or you do not have permission")

        # The taggit manager is smart enough to handle adding/removing based on the list.
        # However, we must ensure that any NEW tags are created for the correct organization.
        # The safest way is to handle this manually.

        tags_for_playbook = []
        for tag_name in tag_names:
            # Get or create the tag specifically for the user's organization
            tag, _ = TenantTag.objects.get_or_create(
                name=tag_name,
                organization=user.organization
            )
            tags_for_playbook.append(tag)

        # The .set() method clears all existing tags and adds the ones in the list.
        playbook.tags.set(tags_for_playbook)
        playbook.save()

        return UpdatePlaybookTags(playbook=playbook)

class CreateTask(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)
        title = graphene.String(required=True)
        description = graphene.String(required=False)
        assignee_id = graphene.ID()
        due_date = graphene.Date()

    task = graphene.Field(TaskType)

    class Meta:
        description = "Creates a new task and associates it with a playbook."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        playbook_id = kwargs.get('playbook_id')
        title = kwargs.get('title')
        description = kwargs.get('description')

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the playbook and verify it belongs to the user's org
        try:
            # Verify Graph ownership
            graph = PlaybookGraph.objects.get(pk=playbook_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Playbook Graph not found or you do not have permission")

        task = Task(playbook=graph, title=title, description=description, **kwargs)
        task.save()
        return CreateTask(task=task)

class UpdateTask(graphene.Mutation):
    class Arguments:
        task_id = graphene.ID(required=True)
        title = graphene.String()
        description = graphene.String()
        status = graphene.String()
        assignee_id = graphene.ID()
        due_date = graphene.Date()

    task = graphene.Field(TaskType)

    class Meta:
        description = "Updates an existing task."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        task_id = kwargs.get('task_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the task and verify its parent playbook belongs to the user's org
        try:
            task = Task.objects.get(pk=task_id, playbook__organization=user.organization)
        except Task.DoesNotExist:
            raise Exception("Task not found or you do not have permission")

        # Update the fields that were provided in the mutation
        for field, value in kwargs.items():
            setattr(task, field, value)

        task.save()
        return UpdateTask(task=task)

class DeleteTask(graphene.Mutation):
    class Arguments:
        task_id = graphene.ID(required=True)

    ok = graphene.Boolean()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        task_id = kwargs.get('task_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the task and verify its parent playbook belongs to the user's org
        try:
            task = Task.objects.get(pk=task_id, playbook__organization=user.organization)
        except Task.DoesNotExist:
            raise Exception("Task not found or you do not have permission")

        task.delete()
        return DeleteTask(ok=True)

# Add this class before the 'class Mutation(...)' line in backend/playbooks/schema.py
class UpdatePlaybookStatus(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)
        status = graphene.String(required=True)

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Updates the status of a single playbook. (Used by Kanban board)"

    @staticmethod
    @role_required([Roles.ADMIN, Roles.REVIEWER])
    def mutate(root, info, id, status, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the playbook and verify it belongs to the user's org
        try:
            playbook = DetectionPlaybook.objects.get(pk=id, organization=user.organization)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission")

        # Validate if the provided status is a valid choice
        valid_statuses = [choice[0] for choice in DetectionPlaybook.PlaybookStatus.choices]
        if status not in valid_statuses:
            raise Exception(f"Invalid status: {status}")

        playbook.status = status
        playbook.save(update_fields=['status', 'updated_at'])  # Optimize DB query

        return UpdatePlaybookStatus(playbook=playbook)

class UpdatePlaybookGraphStatus(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)
        status = graphene.String(required=True)

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Updates the status of a single v2 PlaybookGraph. (Used by Kanban board)"

    @staticmethod
    @role_required([Roles.ADMIN, Roles.REVIEWER])
    def mutate(root, info, id, status, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security: Only allow updates within user's organization
        try:
            graph = PlaybookGraph.objects.get(pk=id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        valid_statuses = [choice[0] for choice in DetectionPlaybook.PlaybookStatus.choices]
        if status not in valid_statuses:
            raise Exception(f"Invalid status: {status}")

        graph.status = status
        graph.save(update_fields=["status", "updated_at"])

        # Create Activity Log
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action=f"Status changed to {status}",
            details=f"User {user.username} changed status to {status}"
        )

        # Create a notification for the author on workbench update (if actor is not the author)
        try:
            if getattr(graph.author, 'id', None) and graph.author != user:
                from django.contrib.contenttypes.models import ContentType
                from notifications.models import Notification
                ct = ContentType.objects.get_for_model(PlaybookGraph)
                Notification.objects.create(
                    recipient=graph.author,
                    actor=user,
                    organization=user.organization,
                    verb=f"Status updated to {status}",
                    object_id=str(graph.id),
                    content_type=ct,
                )
                # Email dispatch respecting preferences
                from core.email_service import get_email_service
                if graph.author and getattr(graph.author, 'email_notify_workbench_edited', False) and graph.author.email:
                    service = get_email_service()
                    if service.is_configured():
                        service.send_message(
                            to=[graph.author.email],
                            subject=f'📊 Workbench Status Changed - {graph.title}',
                            text=f"""Hello {graph.author.username},

Your workbench "{graph.title}" status has been updated.

New Status: {status}
Changed by: {user.username}

Best regards,
The HEFAISTOS Team""",
                            html=f"""<html><body>
<h2>📊 Workbench Status Updated</h2>
<p>Hello <strong>{graph.author.username}</strong>,</p>
<p>Your workbench "<strong>{graph.title}</strong>" status has been updated.</p>
<ul>
<li><strong>New Status:</strong> {status}</li>
<li><strong>Changed by:</strong> {user.username}</li>
</ul>
<p>Best regards,<br/>The HEFAISTOS Team</p>
</body></html>"""
                        )
        except Exception:
            pass

        # Publish status change event for notifications
        try:
            publisher = get_publisher()
            publisher.publish_message('playbook.graph.status.changed', {
                'graph_id': str(graph.id),
                'status': status,
                'organization_id': str(user.organization.id),
                'actor_id': str(user.id),
                'creator_id': str(getattr(graph.author, 'id', '')),
            })
        except Exception:
            pass


        return UpdatePlaybookGraphStatus(playbook_graph=graph)

class UpdateOwnPlaybookGraphStatus(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)
        status = graphene.String(required=True)

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Owner-only: Updates status for a PlaybookGraph with restricted transitions."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.REVIEWER, Roles.VIEWER])
    def mutate(root, info, id, status, **kwargs):
        status = (status or '').upper()

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        try:
            graph = PlaybookGraph.objects.select_related('author').get(pk=id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        # Owner-only
        if getattr(graph.author, 'id', None) != getattr(user, 'id', None):
            raise Exception("Only the owner can change this status here.")

        # Cannot change while in IDEA
        if graph.status == DetectionPlaybook.PlaybookStatus.IDEA:
            raise Exception("Status cannot be changed while in IDEA.")

        allowed = {
            DetectionPlaybook.PlaybookStatus.DEVELOPMENT,
            DetectionPlaybook.PlaybookStatus.TESTING,
            DetectionPlaybook.PlaybookStatus.TUNING,
        }
        if status not in allowed:
            raise Exception("Invalid status. Allowed: DEVELOPMENT, TESTING, TUNING")

        graph.status = status
        graph.save(update_fields=["status", "updated_at"])

        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action=f"Status changed to {status}",
            details=f"Owner changed status to {status}"
        )

        # Notify author (no-op if self-change) and email if opted-in
        try:
            if getattr(graph.author, 'id', None) and graph.author != user:
                from django.contrib.contenttypes.models import ContentType
                from notifications.models import Notification
                ct = ContentType.objects.get_for_model(PlaybookGraph)
                Notification.objects.create(
                    recipient=graph.author,
                    actor=user,
                    organization=user.organization,
                    verb=f"Owner updated status to {status}",
                    object_id=str(graph.id),
                    content_type=ct,
                )
                from core.email_service import get_email_service
                if graph.author and getattr(graph.author, 'email_notify_workbench_edited', False) and graph.author.email:
                    service = get_email_service()
                    if service.is_configured():
                        service.send_message(
                            to=[graph.author.email],
                            subject=f'📊 Workbench Status Changed - {graph.title}',
                            text=f"""Hello {graph.author.username},

Your workbench "{graph.title}" status has been updated by owner.

New Status: {status}
Changed by: {user.username}

Best regards,
The HEFAISTOS Team""",
                            html=f"""<html><body>
<h2>📊 Workbench Status Updated</h2>
<p>Hello <strong>{graph.author.username}</strong>,</p>
<p>Your workbench "<strong>{graph.title}</strong>" status has been updated by owner.</p>
<ul>
<li><strong>New Status:</strong> {status}</li>
<li><strong>Changed by:</strong> {user.username}</li>
</ul>
<p>Best regards,<br/>The HEFAISTOS Team</p>
</body></html>"""
                        )
        except Exception:
            pass

        return UpdateOwnPlaybookGraphStatus(playbook_graph=graph)

class UpdatePlaybookGraphTitle(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)
        title = graphene.String(required=True)

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Owner-only: Renames a PlaybookGraph."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.REVIEWER, Roles.VIEWER])
    def mutate(root, info, id, title, **kwargs):
        title = (title or '').strip()
        if not title:
            raise Exception("Title cannot be empty")

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        try:
            graph = PlaybookGraph.objects.select_related('author').get(pk=id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        # Owner-only
        if getattr(graph.author, 'id', None) != getattr(user, 'id', None):
            raise Exception("Only the owner can rename this workbench.")

        old = graph.title
        graph.title = title[:255]
        graph.save(update_fields=["title", "updated_at"])

        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="Renamed",
            details=f"Title: '{old}' -> '{graph.title}'"
        )

        # Notify author about rename (if actor != author)
        try:
            if getattr(graph.author, 'id', None) and graph.author != user:
                from django.contrib.contenttypes.models import ContentType
                from notifications.models import Notification
                ct = ContentType.objects.get_for_model(PlaybookGraph)
                Notification.objects.create(
                    recipient=graph.author,
                    actor=user,
                    organization=user.organization,
                    verb=f"Renamed workbench to '{graph.title}'",
                    object_id=str(graph.id),
                    content_type=ct,
                )
                from core.email_service import get_email_service
                if graph.author and getattr(graph.author, 'email_notify_workbench_edited', False) and graph.author.email:
                    service = get_email_service()
                    if service.is_configured():
                        service.send_message(
                            to=[graph.author.email],
                            subject=f'✏️ Workbench Renamed - {graph.title}',
                            text=f"""Hello {graph.author.username},

Your workbench has been renamed.

Previous Title: {old}
New Title: {graph.title}
Renamed by: {user.username}

Best regards,
The HEFAISTOS Team""",
                            html=f"""<html><body>
<h2>✏️ Workbench Renamed</h2>
<p>Hello <strong>{graph.author.username}</strong>,</p>
<p>Your workbench has been renamed.</p>
<ul>
<li><strong>Previous Title:</strong> {old}</li>
<li><strong>New Title:</strong> {graph.title}</li>
<li><strong>Renamed by:</strong> {user.username}</li>
</ul>
<p>Best regards,<br/>The HEFAISTOS Team</p>
</body></html>"""
                        )
        except Exception:
            pass

        return UpdatePlaybookGraphTitle(playbook_graph=graph)

class AdminApproveDeployment(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Admin-only: set PlaybookGraph status to DEPLOYED."

    @staticmethod
    @role_required([Roles.ADMIN])
    def mutate(root, info, id):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")
        try:
            graph = PlaybookGraph.objects.get(pk=id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        # Only allow deployment from APPROVED
        if (graph.status or '').upper() != str(DetectionPlaybook.PlaybookStatus.APPROVED):
            raise Exception("Graph must be APPROVED before deployment")

        graph.status = DetectionPlaybook.PlaybookStatus.DEPLOYED
        graph.save(update_fields=["status", "updated_at"])

        # Publish deployment event via existing connector key
        try:
            publisher = get_publisher()
            publisher.publish_message('playbook.graph.status.changed', {
                'graph_id': str(graph.id),
                'status': str(graph.status),
                'organization_id': str(user.organization.id),
                'actor_id': str(user.id),
                'creator_id': str(getattr(graph.author, 'id', '')),
            })
        except Exception:
            pass

        try:
            _queue_dac_deployment_automation(graph, user)
        except Exception as automation_exc:
            export_logger.exception(
                'DaC automation failed for graph %s after DEPLOYED transition',
                graph.id,
            )
            _notify_dac_automation_failure(graph, user, str(automation_exc))


        return AdminApproveDeployment(playbook_graph=graph)

class UpdatePlaybook(graphene.Mutation):
    class Arguments:
        id = graphene.UUID(required=True)
        # --- All fields from CreatePlaybook ---
        title = graphene.String()
        description = graphene.String()
        playbook_type = graphene.String()
        analytic_id = graphene.String(description="Unique ID for the analytic")
        version = graphene.String(description="Semantic version (e.g., 1.0)")
        hypothesis = graphene.String(description="The testable question for a HUNT playbook")
        robustness_level = graphene.Int(description="Level 1-5 of how easy this detection is to evade")
        data_source_robustness = graphene.String(description="Where does the data come from? (e.g., K for Kernel-Mode)")
        false_positive_rate = graphene.Int(description="Estimated FP rate 0-100%")
        known_false_positives = graphene.String(description="Known benign triggers for a DETECTION playbook")
        exclusion_strategy = graphene.String(description="Surgically precise exclusions")
        testing_procedures = graphene.String(description="Test cases and synonym tools")
        triage_guidance = graphene.String(description="Step-by-step instructions for T1 analysts")
        soar_enrichment = graphene.String(description="Automated enrichment steps")
        soar_triage = graphene.String(description="Automated triage logic")
        soar_containment = graphene.String(description="Automated containment steps")
        mitre_attack_mappings = graphene.List(graphene.ID)
        mitre_d3fend_mappings = graphene.List(graphene.ID)
        mitre_engage_mappings = graphene.List(graphene.ID)
        mitre_ics_mappings = graphene.List(graphene.ID)
        mitre_mobile_mappings = graphene.List(graphene.ID)
        operational_path = graphene.String(description="Operational path phases")
        function_call_graphs = graphene.String(description="Function call graph abstractions")
        execution_modalities = graphene.String(description="Execution modalities / behavioral variants")
        technical_details = graphene.String(description="Deep technical implementation notes (HUNT)")

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Updates the core metadata of a playbook."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        id = kwargs.get('id')

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # Security Check: Get the playbook and verify it belongs to the user's org
        try:
            playbook = DetectionPlaybook.objects.get(pk=id, organization=user.organization)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission")

        # --- Handle M2M fields ---
        attack_ids = kwargs.pop('mitre_attack_mappings', None)
        d3fend_ids = kwargs.pop('mitre_d3fend_mappings', None)
        engage_ids = kwargs.pop('mitre_engage_mappings', None)
        ics_ids = kwargs.pop('mitre_ics_mappings', None)
        mobile_ids = kwargs.pop('mitre_mobile_mappings', None)

        # Update simple fields
        for field, value in kwargs.items():
            setattr(playbook, field, value)

        playbook.save()  # Save simple fields

        # --- Set M2M relationships ---
        if attack_ids is not None:
            playbook.mitre_attack_mappings.set(attack_ids)
        if d3fend_ids is not None:
            playbook.mitre_d3fend_mappings.set(d3fend_ids)
        if engage_ids is not None:
            playbook.mitre_engage_mappings.set(engage_ids)
        if ics_ids is not None:
            playbook.mitre_ics_mappings.set(ics_ids)
        if mobile_ids is not None:
            playbook.mitre_mobile_mappings.set(mobile_ids)

        return UpdatePlaybook(playbook=playbook)


class UpdatePlaybookGraphMetadata(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        notes = graphene.String()
        playbook_ids = graphene.List(graphene.UUID, description="DetectionPlaybooks to attach this graph to.")

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Updates graph notes and which detection playbooks it is attached to."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        notes = kwargs.get('notes')
        playbook_ids = kwargs.get('playbook_ids')

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        if notes is not None:
            graph.notes = notes

        if playbook_ids is not None:
            qs = DetectionPlaybook.objects.filter(pk__in=playbook_ids, organization=user.organization)
            graph.playbooks.set(qs)

        graph.save()

        # Create Activity Log
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="Metadata updated",
            details=f"User {user.username} updated metadata"
        )

        # Notify author (actor != author)
        try:
            if getattr(graph.author, 'id', None) and graph.author != user:
                from django.contrib.contenttypes.models import ContentType
                from notifications.models import Notification
                ct = ContentType.objects.get_for_model(PlaybookGraph)
                Notification.objects.create(
                    recipient=graph.author,
                    actor=user,
                    organization=user.organization,
                    verb="Updated workbench metadata",
                    object_id=str(graph.id),
                    content_type=ct,
                )
                from core.email_service import get_email_service
                if graph.author and getattr(graph.author, 'email_notify_workbench_edited', False) and graph.author.email:
                    service = get_email_service()
                    if service.is_configured():
                        service.send_message(
                            to=[graph.author.email],
                            subject=f'📝 Workbench Metadata Updated - {graph.title}',
                            text=f"""Hello {graph.author.username},

Your workbench "{graph.title}" metadata has been updated.

Updated by: {user.username}

Best regards,
The HEFAISTOS Team""",
                            html=f"""<html><body>
<h2>📝 Workbench Metadata Updated</h2>
<p>Hello <strong>{graph.author.username}</strong>,</p>
<p>Your workbench "<strong>{graph.title}</strong>" metadata has been updated.</p>
<p><strong>Updated by:</strong> {user.username}</p>
<p>Best regards,<br/>The HEFAISTOS Team</p>
</body></html>"""
                        )
        except Exception:
            pass

        return UpdatePlaybookGraphMetadata(playbook_graph=graph)


class CreateCapabilityAbstraction(graphene.Mutation):
    class Arguments:
        technique_id = graphene.String(required=True)
        abstraction_layer = graphene.String(required=True)
        component_artifact = graphene.String(required=True)
        adversary_purpose = graphene.String()
        common_evasions = graphene.String()
        expected_observables = graphene.String()
        applicable_telemetry = graphene.String()
        detection_value = graphene.String()
        robustness_level = graphene.Int()
        review_status = graphene.String()

    capability_abstraction = graphene.Field(CapabilityAbstractionType)

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, technique_id, abstraction_layer, component_artifact, **kwargs):
        user = info.context.user
        try:
            technique = MitreAttackTechnique.objects.get(technique_id=technique_id)
        except MitreAttackTechnique.DoesNotExist:
            raise Exception("Technique not found")

        layer = (abstraction_layer or '').upper()
        valid_layers = {choice[0] for choice in CapabilityAbstraction.AbstractionLayer.choices}
        if layer not in valid_layers:
            raise Exception("Invalid abstraction layer")

        review_status = (kwargs.get('review_status') or CapabilityAbstraction.ReviewStatus.DRAFT).upper()
        valid_review_statuses = {choice[0] for choice in CapabilityAbstraction.ReviewStatus.choices}
        if review_status not in valid_review_statuses:
            raise Exception("Invalid review status")

        capability = CapabilityAbstraction.objects.create(
            technique=technique,
            organization=user.organization,
            created_by=user,
            updated_by=user,
            abstraction_layer=layer,
            component_artifact=component_artifact.strip(),
            adversary_purpose=kwargs.get('adversary_purpose') or '',
            common_evasions=kwargs.get('common_evasions') or '',
            expected_observables=kwargs.get('expected_observables') or '',
            applicable_telemetry=kwargs.get('applicable_telemetry') or '',
            detection_value=kwargs.get('detection_value') or '',
            robustness_level=kwargs.get('robustness_level') or 0,
            review_status=review_status,
            source_kind=CapabilityAbstraction.SourceKind.CUSTOM,
            is_baseline=False,
        )
        return CreateCapabilityAbstraction(capability_abstraction=capability)


class UpdateCapabilityAbstraction(graphene.Mutation):
    class Arguments:
        capability_abstraction_id = graphene.UUID(required=True)
        abstraction_layer = graphene.String()
        component_artifact = graphene.String()
        adversary_purpose = graphene.String()
        common_evasions = graphene.String()
        expected_observables = graphene.String()
        applicable_telemetry = graphene.String()
        detection_value = graphene.String()
        robustness_level = graphene.Int()
        review_status = graphene.String()

    capability_abstraction = graphene.Field(CapabilityAbstractionType)

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, capability_abstraction_id, **kwargs):
        user = info.context.user
        try:
            capability = CapabilityAbstraction.objects.get(pk=capability_abstraction_id)
        except CapabilityAbstraction.DoesNotExist:
            raise Exception("Capability abstraction not found")

        if capability.organization_id != getattr(user.organization, 'id', None):
            raise Exception("Only organization-scoped capability abstractions can be edited")

        valid_layers = {choice[0] for choice in CapabilityAbstraction.AbstractionLayer.choices}
        valid_review_statuses = {choice[0] for choice in CapabilityAbstraction.ReviewStatus.choices}
        dirty = False

        if 'abstraction_layer' in kwargs and kwargs['abstraction_layer'] is not None:
            layer = kwargs['abstraction_layer'].upper()
            if layer not in valid_layers:
                raise Exception("Invalid abstraction layer")
            capability.abstraction_layer = layer
            dirty = True
        if 'component_artifact' in kwargs and kwargs['component_artifact'] is not None:
            capability.component_artifact = kwargs['component_artifact'].strip()
            dirty = True
        for field in (
            'adversary_purpose',
            'common_evasions',
            'expected_observables',
            'applicable_telemetry',
            'detection_value',
            'robustness_level',
        ):
            if field in kwargs and kwargs[field] is not None:
                setattr(capability, field, kwargs[field])
                dirty = True
        if 'review_status' in kwargs and kwargs['review_status'] is not None:
            review_status = kwargs['review_status'].upper()
            if review_status not in valid_review_statuses:
                raise Exception("Invalid review status")
            capability.review_status = review_status
            dirty = True

        if dirty:
            capability.version += 1
            capability.updated_by = user
            capability.save()

        return UpdateCapabilityAbstraction(capability_abstraction=capability)


class DeleteCapabilityAbstraction(graphene.Mutation):
    class Arguments:
        capability_abstraction_id = graphene.UUID(required=True)

    ok = graphene.Boolean()

    @staticmethod
    @role_required([Roles.ADMIN])
    def mutate(root, info, capability_abstraction_id):
        user = info.context.user
        try:
            capability = CapabilityAbstraction.objects.get(pk=capability_abstraction_id)
        except CapabilityAbstraction.DoesNotExist:
            raise Exception("Capability abstraction not found")

        if capability.organization_id != getattr(user.organization, 'id', None):
            raise Exception("Only organization-scoped capability abstractions can be deleted")
        if capability.source_kind != CapabilityAbstraction.SourceKind.CUSTOM or capability.is_baseline:
            raise Exception("Only custom capability abstractions can be deleted")

        capability.delete()
        return DeleteCapabilityAbstraction(ok=True)


class SetPlaybookGraphSnapshot(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        png_base64 = graphene.String(required=True, description="Base64-encoded PNG data (data URL or raw).")

    playbook_graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Sets or updates the PNG snapshot for a graph."

    @staticmethod
    def mutate(root, info, graph_id, png_base64, **kwargs):
        import base64
        from django.core.files.base import ContentFile

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        if "," in png_base64:
            _, data = png_base64.split(",", 1)
        else:
            data = png_base64

        decoded = base64.b64decode(data)
        filename = f"graph_{graph.id}.png"
        graph.png_snapshot.save(filename, ContentFile(decoded), save=True)

        return SetPlaybookGraphSnapshot(playbook_graph=graph)

class DeployPlaybook(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Moves an APPROVED playbook to TESTING. This is the hook for RabbitMQ."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        playbook_id = kwargs.get('playbook_id')

        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        try:
            playbook = DetectionPlaybook.objects.get(
                pk=playbook_id, 
                organization=user.organization
            )
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission")

        if playbook.status != DetectionPlaybook.PlaybookStatus.APPROVED:
            raise Exception(f"Playbook must be in APPROVED status to deploy, not {playbook.status}")

        # --- START OF REFACTOR ---
        #
        # REMOVED:
        # playbook.status = DetectionPlaybook.PlaybookStatus.TESTING
        # playbook.save(update_fields=['status', 'updated_at'])
        #
        # ADDED:

        # 1. Define the routing key
        routing_key = "playbook.deploy.testing"

        # 2. Define the message payload
        message_body = {
            "action": "deploy_to_testing",
            "playbook_id": str(playbook.id),
            "organization_id": str(playbook.organization_id),
            "triggered_by_user_id": str(user.id)
        }

        # 3. Get the publisher instance and publish
        try:
            publisher = get_publisher()
            publisher.publish_message(routing_key, message_body)
        except Exception as e:
            # If publishing fails, we must stop and inform the user.
            raise Exception(f"Could not publish deploy event: {e}")

        # --- END OF REFACTOR ---

        # We no longer change the status. We just return the playbook
        # in its current "APPROVED" state. A separate listener
        # (built in Sprint 13) will be responsible for changing the status.
        return DeployPlaybook(playbook=playbook)

# --- NEW: Delete Graph & Playbook (Top-level) ---
class DeletePlaybookGraph(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)

    ok = graphene.Boolean()

    class Meta:
        description = "Deletes a playbook graph and all its nodes/edges (cascade). Only the author can delete, and only if not DEPLOYED."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")
        try:
            graph_filters = {"pk": graph_id}
            if not user.is_superuser:
                graph_filters["organization"] = user.organization
            graph = PlaybookGraph.objects.get(**graph_filters)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")
        
        # Only the author can delete
        if not user.is_superuser and getattr(graph, 'author_id', None) != getattr(user, 'id', None):
            raise Exception("Only the author can delete this workbench.")
        
        # Cannot delete if DEPLOYED
        if (graph.status or '').upper() == 'DEPLOYED':
            raise Exception("Cannot delete a deployed workbench. Please undeploy first.")
        
        # Also delete any associated DetectionPlaybooks that were created for this graph
        # (e.g., when submitting for review). Only delete playbooks that are exclusively
        # linked to this graph (not shared with other graphs).
        linked_playbooks = list(graph.playbooks.all())
        for playbook in linked_playbooks:
            # Check if this playbook is linked to other graphs
            other_graphs_count = playbook.graphs.exclude(pk=graph_id).count()
            if other_graphs_count == 0:
                # This playbook is only linked to this graph, safe to delete
                # Also delete any associated review requests
                from review.models import ReviewRequest as CanonReviewRequest
                CanonReviewRequest.objects.filter(playbook=playbook).delete()
                playbook.delete()
        
        graph.delete()
        return DeletePlaybookGraph(ok=True)

class DeleteDetectionPlaybook(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)

    ok = graphene.Boolean()

    class Meta:
        description = "Deletes a detection or hunt playbook (if belongs to user's org)."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        playbook_id = kwargs.get('playbook_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")
        try:
            playbook_filters = {"pk": playbook_id}
            if not user.is_superuser:
                playbook_filters["organization"] = user.organization
            pb = DetectionPlaybook.objects.get(**playbook_filters)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission")
        pb.delete()
        return DeleteDetectionPlaybook(ok=True)

# --- NEW V2 MUTATIONS ---

class CreatePlaybookGraph(graphene.Mutation):
    class Arguments:
        title = graphene.String(required=True)

    graph = graphene.Field(PlaybookGraphType)

    class Meta:
        description = "Creates a new, blank playbook graph (v2)."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        title = kwargs.get('title')
        user = info.context.user
        graph = PlaybookGraph.objects.create(
            title=title,
            organization=user.organization,
            author=user,
            status='IDEA'
        )

        # Create Activity Log
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="Created",
            details=f"User {user.username} created the playbook"
        )

        # OPTIONAL: Auto-create a default "Root" node
        PlaybookNode.objects.create(
            graph=graph,
            layer_name="Incident Start",
            position_x=100,
            position_y=100
        )

        return CreatePlaybookGraph(graph=graph)

class CreatePlaybookNode(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        layer_name = graphene.String(required=False, default_value="New Node")
        position_x = graphene.Float(required=True)
        position_y = graphene.Float(required=True)

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Adds a new node to an existing playbook graph."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        position_x = kwargs.get('position_x')
        position_y = kwargs.get('position_y')
        layer_name = kwargs.get('layer_name')
        user = info.context.user

        try:
            # Security: Ensure the graph belongs to the user's org
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        node = PlaybookNode.objects.create(
            graph=graph,
            layer_name=layer_name,
            position_x=position_x,
            position_y=position_y
        )
        return CreatePlaybookNode(node=node)

class CreatePlaybookEdge(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        source_id = graphene.UUID(required=True)
        target_id = graphene.UUID(required=True)

    edge = graphene.Field(PlaybookEdgeType)

    class Meta:
        description = "Connects two nodes on a playbook graph."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        source_id = kwargs.get('source_id')
        target_id = kwargs.get('target_id')
        # Simplified logic: In production, check circular dependency or duplicates
        edge = PlaybookEdge.objects.create(
            graph_id=graph_id,
            source_node_id=source_id,
            target_node_id=target_id
        )
        return CreatePlaybookEdge(edge=edge)

class UpdatePlaybookNodePosition(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)
        position_x = graphene.Float(required=True)
        position_y = graphene.Float(required=True)

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Updates the (x, y) position of a node (for drag-and-drop)."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, node_id, position_x, position_y, **kwargs):
        user = info.context.user
        try:
            # Security: Find the node and check that its graph belongs to the user's org
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        node.position_x = position_x
        node.position_y = position_y
        node.save(update_fields=['position_x', 'position_y'])

        return UpdatePlaybookNodePosition(node=node)

class DeletePlaybookNode(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)

    ok = graphene.Boolean()

    class Meta:
        description = "Deletes a node from a graph. (Will also delete connected edges)"

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        node_id = kwargs.get('node_id')
        user = info.context.user
        try:
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        node.delete() # 'on_delete=models.CASCADE' on PlaybookEdge will handle edges
        return DeletePlaybookNode(ok=True)

class DeletePlaybookEdge(graphene.Mutation):
    class Arguments:
        edge_id = graphene.UUID(required=True)

    ok = graphene.Boolean()

    class Meta:
        description = "Deletes an edge from a graph."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
         edge_id = kwargs.get('edge_id')
         PlaybookEdge.objects.get(pk=edge_id).delete()
         return DeletePlaybookEdge(ok=True)

# --- NEW: Update Node Template Data ---
class UpdateNodeTemplate(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)
        template_data = GenericScalar(required=True)

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Updates the template_data JSON for a specific node."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, node_id, template_data, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        # Expect object input; reject non-dict types
        if not isinstance(template_data, dict):
            raise Exception("templateData must be an object")

        node.template_data = template_data
        # PlaybookNode does not define an 'updated_at' field, so only update template_data
        node.save(update_fields=["template_data"])

        return UpdateNodeTemplate(node=node)

# --- NEW: Update Node Layer Name (rename) ---
class UpdatePlaybookNodeLayerName(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)
        layer_name = graphene.String(required=True, description="New display name for the node")

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Renames a playbook graph node (updates layer_name)."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, node_id, layer_name, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        if not layer_name.strip():
            raise Exception("Node name cannot be empty")

        try:
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        node.layer_name = layer_name.strip()[:100]
        node.save(update_fields=["layer_name"])
        return UpdatePlaybookNodeLayerName(node=node)

# --- NEW: Update Node ATT&CK Mappings ---
class UpdateNodeAttackMappings(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)
        mitre_attack_ids = graphene.List(graphene.ID, required=True)

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Sets the MITRE ATT&CK mappings for a specific node."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        import uuid as uuid_module
        import logging
        _logger = logging.getLogger(__name__)
        
        node_id = kwargs.get('node_id')
        mitre_attack_ids = kwargs.get('mitre_attack_ids') or []
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        # Accept both UUID PKs and ATT&CK technique IDs (e.g., T1059, T1003.001)
        valid_uuids = []
        technique_ids = []
        for id_str in mitre_attack_ids:
            if not id_str:
                continue
            try:
                uid = uuid_module.UUID(str(id_str))
                valid_uuids.append(uid)
                continue
            except (ValueError, TypeError):
                # Not a UUID, try treating as technique_id (e.g., T1059)
                technique_ids.append(str(id_str).strip())
                continue

        techniques_qs = MitreAttackTechnique.objects.none()
        if valid_uuids:
            techniques_qs = techniques_qs | MitreAttackTechnique.objects.filter(pk__in=valid_uuids)
        if technique_ids:
            techniques_qs = techniques_qs | MitreAttackTechnique.objects.filter(technique_id__in=technique_ids)

        # Deduplicate
        node.mitre_attack_mappings.set(list({t.id: t for t in techniques_qs}.values()))
        node.save()
        return UpdateNodeAttackMappings(node=node)


# --- NEW: Share / Un-share Playbook Across Entity ---
class SharePlaybook(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True)
        share = graphene.Boolean(required=True, description="Set to 'true' to share, 'false' to un-share.")

    playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Shares or un-shares a playbook with the user's entire Entity."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, playbook_id, share, **kwargs):
        user = info.context.user

        # --- Security Check 1: Can only share playbooks YOU own ---
        try:
            playbook = DetectionPlaybook.objects.get(pk=playbook_id, organization=user.organization)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Playbook not found or you do not have permission to share it.")

        # --- Business Logic Check ---
        if not getattr(user, "organization", None) or not getattr(user.organization, "entity", None):
            raise Exception("Your organization is not part of an Entity, so you cannot share playbooks.")

        # --- Action ---
        playbook.is_shared = share
        playbook.save(update_fields=["is_shared", "updated_at"])

        return SharePlaybook(playbook=playbook)


# --- NEW: Share / Un-share PlaybookGraph (Workbench) Across Entity ---
class SharePlaybookGraph(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        share = graphene.Boolean(required=True, description="Set to 'true' to share, 'false' to un-share.")

    graph = graphene.Field(PlaybookGraphType)
    success = graphene.Boolean()
    message = graphene.String()

    class Meta:
        description = "Shares or un-shares a Workbench (PlaybookGraph) with all organizations in the user's Entity."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id, share, **kwargs):
        user = info.context.user
        
        if user.is_anonymous:
            return SharePlaybookGraph(graph=None, success=False, message="Authentication required")

        # Security: Only allow sharing of graphs owned by user's organization
        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            return SharePlaybookGraph(graph=None, success=False, message="Workbench not found or you do not have permission")

        # Only author can share
        if getattr(graph, 'author_id', None) != getattr(user, 'id', None):
            # Allow admins to share any graph in their org
            if user.role != 'ADMIN':
                return SharePlaybookGraph(graph=None, success=False, message="Only the author or an admin can share this workbench")

        # Business Logic: Organization must be part of an Entity to share
        if not getattr(user, "organization", None) or not getattr(user.organization, "entity", None):
            return SharePlaybookGraph(graph=None, success=False, message="Your organization is not part of an Entity. Sharing is only available within Entities.")

        # Perform the share/unshare
        graph.is_shared = share
        graph.save(update_fields=["is_shared", "updated_at"])

        # Log activity
        action = "Shared with Entity" if share else "Unshared from Entity"
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action=action,
            details=f"Workbench {'shared with' if share else 'unshared from'} all organizations in {user.organization.entity.name}"
        )

        return SharePlaybookGraph(
            graph=graph, 
            success=True, 
            message=f"Workbench {'shared' if share else 'unshared'} successfully"
        )


class ClonePlaybookGraph(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)

    playbook_graph = graphene.Field(lambda: PlaybookGraphType)
    ok = graphene.Boolean()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Workbench not found or you do not have permission")

        existing_titles = set(PlaybookGraph.objects.filter(organization=user.organization).values_list('title', flat=True))
        new_title = generate_copy_title(graph.title, existing_titles)

        new_graph = PlaybookGraph.objects.create(
            title=new_title,
            organization=user.organization,
            author=user,
            status=graph.status,
            png_snapshot=graph.png_snapshot,
            mitre_technique=graph.mitre_technique,
            selected_strategy=graph.selected_strategy,
            detection_rule=graph.detection_rule,
            goal=graph.goal,
            technical_context=graph.technical_context,
            blind_spots=graph.blind_spots,
            triage_guidance=graph.triage_guidance,
            false_positives=graph.false_positives,
            response_playbook=graph.response_playbook,
            test_scenario=graph.test_scenario,
            test_expected_output=graph.test_expected_output,
            target_file_path=graph.target_file_path,
            git_status=graph.git_status,
            last_commit_hash=graph.last_commit_hash,
            version=graph.version,
            robustness_level=graph.robustness_level,
            data_source_robustness=graph.data_source_robustness,
            is_shared=False,
            notes=graph.notes,
            alert_trigger=graph.alert_trigger,
            default_severity=graph.default_severity,
            enrichment_steps=graph.enrichment_steps,
            containment_steps=graph.containment_steps,
            notification_steps=graph.notification_steps,
            downstream_correlation_requirements=graph.downstream_correlation_requirements,
        )

        new_graph.tags.set(graph.tags.all())
        new_graph.playbooks.set(graph.playbooks.all())

        for task in graph.tasks.all():
            Task.objects.create(
                playbook=new_graph,
                title=task.title,
                description=task.description,
                status=task.status,
                assignee=task.assignee,
                due_date=task.due_date,
            )

        node_map = {}
        for node in graph.nodes.all():
            clone_node = PlaybookNode.objects.create(
                graph=new_graph,
                layer_name=node.layer_name,
                position_x=node.position_x,
                position_y=node.position_y,
                color=node.color,
                ui_metadata=node.ui_metadata,
            )
            clone_node.mitre_attack_mappings.set(node.mitre_attack_mappings.all())
            node_map[node.id] = clone_node

        for edge in graph.edges.all():
            source = node_map.get(edge.source_node_id)
            target = node_map.get(edge.target_node_id)
            if source and target:
                PlaybookEdge.objects.create(
                    graph=new_graph,
                    source_node=source,
                    target_node=target,
                )

        ActivityLog.objects.create(
            playbook=new_graph,
            user=user,
            action="CLONED",
            details=f"Cloned from {graph.title}"
        )

        return ClonePlaybookGraph(playbook_graph=new_graph, ok=True)


class ClonePlaybook(graphene.Mutation):
    class Arguments:
        playbook_id = graphene.UUID(required=True, description="The ID of the shared playbook to use as a template.")
        new_title = graphene.String(required=True, description="The title for the new, cloned playbook.")

    new_playbook = graphene.Field(PlaybookType)

    class Meta:
        description = "Clones a shared playbook into the user's own organization, creating a new, editable copy."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        playbook_id = kwargs.get('playbook_id')
        new_title = kwargs.get('new_title')
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication credentials were not provided")

        # --- 1. Find the "template" playbook ---
        try:
            template_playbook = DetectionPlaybook.objects.get(pk=playbook_id)
        except DetectionPlaybook.DoesNotExist:
            raise Exception("Template playbook not found")

        # --- 2. Clone the playbook ---
        new_playbook = DetectionPlaybook(
            title=new_title,
            description=template_playbook.description,
            playbook_type=template_playbook.playbook_type,
            author=user,
            organization=user.organization,
            status=DetectionPlaybook.PlaybookStatus.IDEA,  # Start as IDEA
            # Copy over other fields as necessary...
        )
        new_playbook.save()

        # --- 3. Copy M2M relationships (tags, rules, data sources, techniques) ---
        new_playbook.tags.set(template_playbook.tags.all())
        new_playbook.detection_rules.set(template_playbook.detection_rules.all())
        new_playbook.required_data_sources.set(template_playbook.required_data_sources.all())
        new_playbook.mitre_attack_mappings.set(template_playbook.mitre_attack_mappings.all())
        new_playbook.mitre_ics_mappings.set(template_playbook.mitre_ics_mappings.all())
        new_playbook.mitre_mobile_mappings.set(template_playbook.mitre_mobile_mappings.all())

        # --- 4. Copy tasks (if any) ---
        for task in template_playbook.tasks.all():
            Task.objects.create(
                playbook=new_playbook,
                title=task.title,
                description=task.description,
                assignee_id=task.assignee_id,
                due_date=task.due_date,
                status=task.status,  # Copy status as well
            )

        return ClonePlaybook(new_playbook=new_playbook)

class UploadGraphSnapshot(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        file = Upload(required=True) # This handles the binary file

    success = graphene.Boolean()
    image_url = graphene.String()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id, file, **kwargs):
        user = info.context.user
        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)

            # Save the file to the model
            # 'file' here is a Django UploadedFile object
            graph.png_snapshot.save(f"{graph.id}_snapshot.png", file, save=True)

            image_url = info.context.build_absolute_uri(graph.png_snapshot.url)
            # Force HTTPS if the request came over HTTPS
            if info.context.is_secure():
                image_url = image_url.replace('http://', 'https://')

            return UploadGraphSnapshot(
                success=True, 
                image_url=image_url
            )
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found")

# --- NEW: Update Playbook Details (Strategy & Context) ---
class UpdatePlaybookDetails(graphene.Mutation):
    """
    Updates the core strategy and context fields of a PlaybookGraph.
    """
    class Arguments:
        graph_id = graphene.UUID(required=True)
        
        # Strategy
        mitre_technique_id = graphene.String()
        selected_strategy = graphene.JSONString()
        detection_rule = graphene.String()
        
        # Context
        goal = graphene.String()
        technical_context = graphene.String()
        blind_spots = graphene.String()
        triage_guidance = graphene.String()
        false_positives = graphene.String()
        response_playbook = graphene.String()
        target_file_path = graphene.String()
        
        # Valuation
        robustness_level = graphene.Int()
        data_source_robustness = graphene.String()
        
        # Maieutic Engine fields
        data_source_maturity = graphene.String()
        conversation_history = graphene.JSONString()
        selected_capability_abstraction_ids = graphene.List(
            graphene.UUID,
            description="Capability abstractions selected to ground AI generation.",
        )
        detection_focus_layer = graphene.String(description="The abstraction layer AI generation should prioritize.")

        # Testing Guidance
        test_scenario = graphene.String()
        test_expected_output = graphene.String()
        
        # Tags (simple list of tag names)
        tags = graphene.List(graphene.String)
        
        # SOAR Configuration
        alert_trigger = graphene.String()
        default_severity = graphene.String()
        enrichment_steps = graphene.JSONString()
        containment_steps = graphene.JSONString()
        notification_steps = graphene.JSONString()
        downstream_correlation_requirements = graphene.JSONString()

        # OpenTide v2.1 fields
        tlp_classification = graphene.String()
        public_references = graphene.JSONString()
        internal_references = graphene.JSONString()
        threat_actors = graphene.JSONString()
        threat_surface = graphene.JSONString()

        # Data sources (Analyst can add via Use Logic)
        required_data_source_ids = graphene.List(graphene.ID, description="List of DataSource IDs to attach to related Playbooks")
        
        # D3FEND Techniques
        d3fend_technique_ids = graphene.List(graphene.String, description="List of D3FEND technique IDs (e.g., D3-LFP)")

    graph = graphene.Field(PlaybookGraphType)

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, **kwargs):
        graph_id = kwargs.get('graph_id')
        user = info.context.user
        
        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found")

        # 1. Handle Technique Change (Triggers ID Generation)
        if 'mitre_technique_id' in kwargs:
            tid = kwargs.get('mitre_technique_id')
            if tid:
                try:
                    tech = MitreAttackTechnique.objects.get(technique_id=tid)
                    if graph.mitre_technique != tech:
                        graph.mitre_technique = tech
                        # Regenerate ID because technique changed
                        graph.custom_id = None # clear old ID
                        graph.generate_custom_id()
                except MitreAttackTechnique.DoesNotExist:
                    pass # Or raise error
            else:
                graph.mitre_technique = None

        # 2. Handle Version Bump (semantic versioning: major.minor)
        # Major bump (detection_rule or robustness_level change) → increment major, reset minor
        if 'detection_rule' in kwargs or 'robustness_level' in kwargs:
            graph.version += 1
            graph.minor_version = 0
        elif any(k in kwargs for k in (
            'alert_trigger', 'default_severity', 'enrichment_steps',
            'containment_steps', 'notification_steps', 'downstream_correlation_requirements', 'tags',
            'goal', 'technical_context', 'blind_spots', 'triage_guidance',
            'false_positives', 'response_playbook', 'target_file_path',
            'data_source_robustness', 'test_scenario', 'test_expected_output',
            'selected_capability_abstraction_ids', 'detection_focus_layer',
            'tlp_classification', 'public_references', 'internal_references', 'threat_actors', 'threat_surface',
        )):
            graph.minor_version += 1

        # 2a. Server-side sanitization for detection_rule (strip markdown code fences/backticks)
        def _sanitize_code_fences(text: str) -> str:
            import re
            if not isinstance(text, str):
                return text
            # Remove triple-fenced blocks ```lang ... ``` and plain ``` ... ```
            cleaned = re.sub(r"```[a-zA-Z0-9_-]*\n([\s\S]*?)```", lambda m: m.group(1), text)
            cleaned = re.sub(r"```([\s\S]*?)```", lambda m: m.group(1), cleaned)
            # Remove stray single backticks
            cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
            # Normalize CRLF and trim
            cleaned = cleaned.replace("\r\n", "\n").strip()
            return cleaned

        def _inject_tags_comment(rule: str, tags: list) -> str:
            """Inject or replace a tags comment line at the top of the detection rule.

            Uses the correct comment syntax for each rule format:
              - KQL  → ``// tags: ...``
              - WAZUH → ``<!-- tags: ... -->``
              - SPL / OTHER → ``# tags: ...``
            """
            import re
            from rules.utils import detect_rule_format
            fmt = detect_rule_format(rule)

            # Remove any pre-existing tags comment lines regardless of style,
            # so stale comments are cleaned up when the format changes.
            rule = re.sub(r"^# tags:.*$\n?", "", rule, flags=re.MULTILINE)
            rule = re.sub(r"^// tags:.*$\n?", "", rule, flags=re.MULTILINE)
            rule = re.sub(r"^<!--\s*tags:.*?-->\n?", "", rule, flags=re.MULTILINE)
            rule = rule.strip()

            if not tags:
                return rule

            if fmt == 'KQL':
                tag_line = f"// tags: {', '.join(tags)}"
            elif fmt == 'WAZUH':
                tag_line = f"<!-- tags: {', '.join(tags)} -->"
            else:
                # SPL, OTHER
                tag_line = f"# tags: {', '.join(tags)}"

            return tag_line + "\n" + rule if rule else tag_line

        # 3. Update standard fields
        if 'selected_strategy' in kwargs: graph.selected_strategy = kwargs['selected_strategy']
        if 'detection_rule' in kwargs:
            raw_rule = kwargs['detection_rule']
            graph.detection_rule = _sanitize_code_fences(raw_rule) if isinstance(raw_rule, str) else raw_rule
        if 'goal' in kwargs: graph.goal = kwargs['goal']
        if 'technical_context' in kwargs: graph.technical_context = kwargs['technical_context']
        if 'blind_spots' in kwargs: graph.blind_spots = kwargs['blind_spots']
        if 'triage_guidance' in kwargs: graph.triage_guidance = kwargs['triage_guidance']
        if 'false_positives' in kwargs: graph.false_positives = kwargs['false_positives']
        if 'response_playbook' in kwargs: graph.response_playbook = kwargs['response_playbook']
        if 'target_file_path' in kwargs: graph.target_file_path = kwargs['target_file_path']
        if 'robustness_level' in kwargs: graph.robustness_level = kwargs['robustness_level']
        if 'data_source_robustness' in kwargs: graph.data_source_robustness = kwargs['data_source_robustness']
        
        # Maieutic Engine fields
        if 'data_source_maturity' in kwargs: graph.data_source_maturity = kwargs['data_source_maturity']
        if 'conversation_history' in kwargs: graph.conversation_history = kwargs['conversation_history']
        if 'detection_focus_layer' in kwargs:
            focus_layer = kwargs['detection_focus_layer'] or ''
            valid_layers = {choice[0] for choice in CapabilityAbstraction.AbstractionLayer.choices}
            focus_layer = focus_layer.upper() if focus_layer else ''
            if focus_layer and focus_layer not in valid_layers:
                raise Exception("Invalid detection focus layer")
            graph.detection_focus_layer = focus_layer

        # 3a. Update testing guidance fields
        if 'test_scenario' in kwargs: graph.test_scenario = kwargs['test_scenario']
        if 'test_expected_output' in kwargs: graph.test_expected_output = kwargs['test_expected_output']
        
        # Tags: replace existing tags with provided list (Tenant-scoped)
        if 'tags' in kwargs:
            from tags.models import TenantTag
            raw_tags = kwargs.get('tags') or []
            # Normalize: lowercase and unique while preserving order
            seen = set()
            norm = []
            for t in raw_tags:
                if not isinstance(t, str):
                    continue
                name = t.strip().lower()
                if name and name not in seen:
                    seen.add(name)
                    norm.append(name)
            tenant_tags = []
            for name in norm:
                tag, _ = TenantTag.objects.get_or_create(
                    name=name,
                    defaults={
                        'slug': name.replace(' ', '-'),
                        'organization': user.organization,
                    },
                    organization=user.organization
                )
                # If tag exists but belongs to another org (shouldn't with unique_together), skip
                if tag.organization_id != user.organization_id:
                    continue
                tenant_tags.append(tag)
            graph.tags.set(tenant_tags)
            # Inject tags as comments into the detection rule
            current_rule = graph.detection_rule or ""
            graph.detection_rule = _inject_tags_comment(current_rule, norm)

        # SOAR Fields — capture old values first for activity log comparison
        soar_subsection_checks = [
            ("Trigger & Severity", ['alert_trigger', 'default_severity']),
            ("Enrichment Steps (Data Gathering)", ['enrichment_steps']),
            ("Containment (Response Actions)", ['containment_steps']),
            ("Notifications", ['notification_steps']),
            ("Downstream Correlation Requirements", ['downstream_correlation_requirements']),
            ("OpenTide Classification & References", ['tlp_classification', 'public_references', 'internal_references']),
            ("Threat Surface Taxonomy", ['threat_surface']),
            ("Threat Actor Attribution", ['threat_actors']),
        ]
        soar_old_values = {}
        for _, fields in soar_subsection_checks:
            for f in fields:
                if f in kwargs:
                    soar_old_values[f] = getattr(graph, f, None)

        if 'alert_trigger' in kwargs: graph.alert_trigger = kwargs['alert_trigger']
        if 'default_severity' in kwargs: graph.default_severity = kwargs['default_severity']
        if 'enrichment_steps' in kwargs: graph.enrichment_steps = kwargs['enrichment_steps']
        if 'containment_steps' in kwargs: graph.containment_steps = kwargs['containment_steps']
        if 'notification_steps' in kwargs: graph.notification_steps = kwargs['notification_steps']
        if 'downstream_correlation_requirements' in kwargs: graph.downstream_correlation_requirements = kwargs['downstream_correlation_requirements'] or {}

        # OpenTide v2.1 fields
        if 'tlp_classification' in kwargs: graph.tlp_classification = kwargs['tlp_classification']
        if 'public_references' in kwargs: graph.public_references = kwargs['public_references'] or []
        if 'internal_references' in kwargs: graph.internal_references = kwargs['internal_references'] or []
        if 'threat_actors' in kwargs: graph.threat_actors = kwargs['threat_actors'] or []
        if 'threat_surface' in kwargs: graph.threat_surface = kwargs['threat_surface'] or []

        # Determine which SOAR subsections actually changed (compare old vs new)
        soar_changed_sections = []
        for section_name, fields in soar_subsection_checks:
            for f in fields:
                if f in soar_old_values:
                    if soar_old_values[f] != getattr(graph, f, None):
                        soar_changed_sections.append(section_name)
                        break

        # 4. Required Data Sources: update on any linked DetectionPlaybooks
        if 'required_data_source_ids' in kwargs:
            ids = [str(x) for x in (kwargs.get('required_data_source_ids') or [])]
            if ids:
                try:
                    from data_catalog.models import DataSource
                    sources = list(DataSource.objects.filter(id__in=ids, organization=user.organization))
                    # Attach to all playbooks linked to this graph
                    for pb in graph.playbooks.all():
                        pb.required_data_sources.set(sources)
                        pb.save(update_fields=['updated_at']) if hasattr(pb, 'updated_at') else pb.save()
                except Exception:
                    # Do not block update on DS errors
                    pass

        if 'selected_capability_abstraction_ids' in kwargs:
            selected_ids = [str(x) for x in (kwargs.get('selected_capability_abstraction_ids') or [])]
            allowed_capabilities = CapabilityAbstraction.objects.filter(
                id__in=selected_ids,
            ).filter(
                Q(organization=user.organization) | Q(organization__isnull=True)
            )
            graph.selected_capability_abstractions.set(allowed_capabilities)
            if not selected_ids:
                graph.detection_focus_layer = ''
            elif not graph.detection_focus_layer and allowed_capabilities.exists():
                graph.detection_focus_layer = allowed_capabilities.first().abstraction_layer

        # 5. D3FEND Techniques: update on the PlaybookGraph
        if 'd3fend_technique_ids' in kwargs:
            d3fend_ids = kwargs.get('d3fend_technique_ids') or []
            if d3fend_ids:
                try:
                    # Look up D3FEND techniques by d3fend_id
                    techniques = D3fendDefensiveTechnique.objects.filter(d3fend_id__in=d3fend_ids)
                    # Set on the graph's M2M relationship
                    graph.d3fend_techniques.set(techniques)
                except Exception as e:
                    # Non-blocking error handling
                    export_logger.warning(f"Failed to set D3FEND techniques: {str(e)}")
            else:
                # Clear D3FEND techniques if empty list provided
                graph.d3fend_techniques.clear()

        graph.save()

        # Create Activity Log with specific details about what changed
        strategy_context_fields = {
            'detection_rule', 'selected_strategy', 'goal', 'technical_context', 'blind_spots',
            'triage_guidance', 'false_positives', 'response_playbook', 'target_file_path',
            'robustness_level', 'data_source_robustness', 'data_source_maturity',
            'test_scenario', 'test_expected_output', 'tags', 'mitre_technique_id',
            'conversation_history', 'required_data_source_ids', 'd3fend_technique_ids',
            'selected_capability_abstraction_ids', 'detection_focus_layer',
        }
        if soar_changed_sections:
            action = "SOAR configuration updated"
            details = f"User {user.username} updated SOAR configuration: {', '.join(soar_changed_sections)}"
        elif any(k in kwargs for k in strategy_context_fields):
            action = "Details updated"
            details = f"User {user.username} updated playbook details (strategy/context)"
        else:
            action = "Details updated"
            details = f"User {user.username} updated playbook details"
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action=action,
            details=details
        )

        return UpdatePlaybookDetails(graph=graph)

# --- NEW: Update Node Color ---
class UpdatePlaybookNodeColor(graphene.Mutation):
    class Arguments:
        node_id = graphene.UUID(required=True)
        color = graphene.String(required=True, description="Node color: default, blue, green, yellow, red")

    node = graphene.Field(PlaybookNodeType)

    class Meta:
        description = "Updates the color of a playbook graph node."

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, node_id, color, **kwargs):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            node = PlaybookNode.objects.get(pk=node_id, graph__organization=user.organization)
        except PlaybookNode.DoesNotExist:
            raise Exception("Node not found or you do not have permission")

        node.color = color
        node.save(update_fields=["color"])
        return UpdatePlaybookNodeColor(node=node)

class AddPlaybookComment(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        message = graphene.String(required=True)

    comment = graphene.Field(PlaybookCommentType)

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST, Roles.REVIEWER])
    def mutate(root, info, graph_id, message):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")
        try:
            playbook = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
            comment = PlaybookComment.objects.create(
                playbook=playbook,
                user=user,
                message=message
            )
            return AddPlaybookComment(comment=comment)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Playbook not found")

class SubmitForReview(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        note = graphene.String() # Optional opening comment

    success = graphene.Boolean()
    review_request = graphene.Field(ReviewRequestType)
    # Simplify payload to avoid type conflicts

    @staticmethod
    @role_required([Roles.ANALYST, Roles.ADMIN])
    def mutate(root, info, graph_id, note=None):
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"[SubmitForReview] Starting mutation for graph_id={graph_id}, note={note}")
        
        user = info.context.user
        # Get the graph (v2)
        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
            logger.info(f"[SubmitForReview] Found graph: {graph.id}, current status={graph.status}")
        except PlaybookGraph.DoesNotExist:
            raise Exception("Graph not found or you do not have permission")

        # Only the author can submit this workbench for review
        if getattr(graph, 'author_id', None) != getattr(user, 'id', None):
            raise Exception("Only the author can submit this workbench for review.")

        # Bridge to canonical review model which reviews DetectionPlaybook
        # If the graph is not attached to a DetectionPlaybook, create one and attach it.
        playbook = graph.playbooks.first()
        if not playbook:
            logger.info(f"[SubmitForReview] No playbook attached, creating new DetectionPlaybook")
            playbook = DetectionPlaybook.objects.create(
                title=graph.title,
                description="",
                playbook_type=DetectionPlaybook.PlaybookType.HUNT if hasattr(DetectionPlaybook, 'PlaybookType') else 'HUNT',
                author=user,
                organization=user.organization,
                status=DetectionPlaybook.PlaybookStatus.REVIEW if hasattr(DetectionPlaybook, 'PlaybookStatus') else 'REVIEW',
            )
            graph.playbooks.add(playbook)
            graph.status = 'REVIEW'
            graph.save()  # Use simple save() to ensure auto_now fields work correctly
            logger.info(f"[SubmitForReview] Created playbook {playbook.id} and attached to graph")
        else:
            logger.info(f"[SubmitForReview] Using existing playbook: {playbook.id}")

        # 1. Create canonical ReviewRequest (review.models) so it matches ReviewRequestType
        req = CanonReviewRequest.objects.create(
            playbook=playbook,
            author=user,
            organization=user.organization,
            status=CanonReviewRequest.ReviewStatus.OPEN,
        )
        logger.info(f"[SubmitForReview] Created ReviewRequest: {req.id}")

        # 2. Update Playbook and Graph Status to REVIEW for UI consistency
        playbook.status = DetectionPlaybook.PlaybookStatus.REVIEW if hasattr(DetectionPlaybook, 'PlaybookStatus') else 'REVIEW'
        playbook.save()  # Use simple save() to ensure auto_now fields work correctly
        logger.info(f"[SubmitForReview] Updated playbook status to: {playbook.status}")
        
        graph.status = 'REVIEW'
        graph.save()  # Use simple save() to ensure auto_now fields work correctly
        logger.info(f"[SubmitForReview] Updated graph status to: {graph.status}")

        # 3. Optional opening comment
        if note:
            comment = CanonReviewComment.objects.create(review_request=req, author=user, text=note)
            logger.info(f"[SubmitForReview] Created comment: {comment.id}")

        # 4. Notify reviewers (Admin/Reviewer roles) via RabbitMQ
        try:
            payload = {
                'event': 'review.requested',
                'organization_id': str(user.organization_id),
                'graph_id': str(graph.id),
                'playbook_id': str(playbook.id),
                'review_request_id': str(req.id),
                'title': graph.title,
                'actor_id': str(user.id),
                'author_id': str(user.id),
                'author_username': getattr(user, 'username', None),
            }
            publish_event('review.requested', payload)
            logger.info(f"[SubmitForReview] Published review.requested event")
        except Exception as e:
            # Log but do not block submit on notification failures
            logger.warning(f"Failed to publish review.requested event: {e}")

        logger.info(f"[SubmitForReview] Mutation complete, returning success=True")
        return SubmitForReview(success=True, review_request=req)

class FinalizeReview(graphene.Mutation):
    class Arguments:
        request_id = graphene.UUID(required=True)
        decision = graphene.String(required=True) # "APPROVE" or "REJECT"
        feedback = graphene.String(required=True) # Mandatory feedback

    success = graphene.Boolean()

    @staticmethod
    @role_required([Roles.REVIEWER, Roles.ADMIN])
    def mutate(root, info, request_id, decision, feedback):
        user = info.context.user
        # Use canonical review request
        try:
            req = CanonReviewRequest.objects.select_related('playbook').get(pk=request_id, organization=user.organization)
        except CanonReviewRequest.DoesNotExist:
            raise Exception("Review request not found or you do not have permission")

        if req.status not in (CanonReviewRequest.ReviewStatus.OPEN,):
            raise Exception("Review is already finalized.")

        # 1. Add formal feedback
        CanonReviewComment.objects.create(review_request=req, author=user, text=feedback)

        # 2. Finalize
        req.updated_at = timezone.now() if hasattr(req, 'updated_at') else req.updated_at

        # Get all graphs linked to this playbook so we can update their status too
        linked_graphs = list(req.playbook.graphs.all())

        if decision.upper() == 'APPROVE':
            req.status = CanonReviewRequest.ReviewStatus.APPROVED
            req.playbook.status = DetectionPlaybook.PlaybookStatus.APPROVED if hasattr(DetectionPlaybook, 'PlaybookStatus') else 'APPROVED'
            new_graph_status = 'APPROVED'
            # Publish review.approved with full details for notifications
            publish_event('review.approved', {
                'playbook_id': str(req.playbook.id),
                'review_request_id': str(req.id),
                'organization_id': str(user.organization_id),
                'actor_id': str(user.id),
                'author_id': str(req.author_id) if req.author_id else None,
                'decision': 'APPROVED',
            })
        elif decision.upper() == 'REJECT':
            req.status = CanonReviewRequest.ReviewStatus.CHANGES_REQUESTED
            req.playbook.status = DetectionPlaybook.PlaybookStatus.RESEARCH if hasattr(DetectionPlaybook, 'PlaybookStatus') else 'RESEARCH'
            new_graph_status = 'RESEARCH'
            # Publish review.changes_requested with full details for notifications
            publish_event('review.changes_requested', {
                'playbook_id': str(req.playbook.id),
                'review_request_id': str(req.id),
                'organization_id': str(user.organization_id),
                'actor_id': str(user.id),
                'author_id': str(req.author_id) if req.author_id else None,
                'decision': 'CHANGES_REQUESTED',
            })
        else:
            raise Exception("Decision must be APPROVE or REJECT")

        req.save(update_fields=['status'])
        req.playbook.save()  # Use simple save() for auto_now fields

        # Update all linked PlaybookGraphs to match the new status
        for graph in linked_graphs:
            graph.status = new_graph_status
            graph.save()  # Use simple save() for auto_now fields
            # Broadcast graph status change for UI caches/notifications
            try:
                publisher = get_publisher()
                publisher.publish_message('playbook.graph.status.changed', {
                    'graph_id': str(graph.id),
                    'status': new_graph_status,
                    'organization_id': str(user.organization.id),
                    'actor_id': str(user.id),
                    'creator_id': str(getattr(graph.author, 'id', '')),
                })
            except Exception:
                pass

        return FinalizeReview(success=True)


# =============================================================================
# EXPORT / IMPORT PLAYBOOK FEATURE
# =============================================================================

def serialize_playbook_graph(graph: PlaybookGraph) -> dict:
    """
    Serialize a PlaybookGraph to a dictionary for export.
    Includes all nodes, edges, and configuration data.
    """
    # Serialize nodes
    nodes = []
    for node in graph.nodes.all():
        node_data = {
            "id": str(node.id),
            "layer_name": node.layer_name,
            "position_x": node.position_x,
            "position_y": node.position_y,
            "color": node.color,
            "ui_metadata": node.ui_metadata or {},
            "mitre_attack_ids": [t.technique_id for t in node.mitre_attack_mappings.all()]
        }
        nodes.append(node_data)
    
    # Serialize edges
    edges = []
    for edge in graph.edges.all():
        edge_data = {
            "id": str(edge.id),
            "source_node_id": str(edge.source_node_id),
            "target_node_id": str(edge.target_node_id)
        }
        edges.append(edge_data)
    
    # Build export object
    export_data = {
        "hefaistos_version": "1.0",
        "export_type": "playbook_graph",
        "exported_at": timezone.now().isoformat(),
        "playbook": {
            # Identity
            "title": graph.title,
            "custom_id": graph.custom_id,
            "version": graph.version,
            "status": graph.status,
            "tags": list(graph.tags.names()),
            
            # Strategy
            "mitre_technique_id": graph.mitre_technique.technique_id if graph.mitre_technique else None,
            "selected_strategy": graph.selected_strategy or {},
            "detection_rule": graph.detection_rule or "",
            
            # Context / Deep Dive
            "goal": graph.goal or "",
            "technical_context": graph.technical_context or "",
            "blind_spots": graph.blind_spots or "",
            "triage_guidance": graph.triage_guidance or "",
            "false_positives": graph.false_positives or "",
            "response_playbook": graph.response_playbook or "",
            
            # Testing
            "test_scenario": graph.test_scenario or "",
            "test_expected_output": graph.test_expected_output or "",
            
            # Deployment
            "target_file_path": graph.target_file_path or "",
            
            # Valuation
            "robustness_level": graph.robustness_level,
            "data_source_robustness": graph.data_source_robustness or "",
            
            # SOAR Configuration
            "alert_trigger": graph.alert_trigger or "",
            "default_severity": graph.default_severity or "MEDIUM",
            "enrichment_steps": graph.enrichment_steps or [],
            "containment_steps": graph.containment_steps or [],
            "notification_steps": graph.notification_steps or [],
            "downstream_correlation_requirements": graph.downstream_correlation_requirements or {},
            
            # Notes
            "notes": graph.notes or "",
            
            # Graph structure
            "nodes": nodes,
            "edges": edges
        }
    }
    
    return export_data


def deserialize_playbook_graph(data: dict, organization, author) -> PlaybookGraph:
    """
    Create a new PlaybookGraph from imported data.
    Returns the created graph.
    """
    playbook_data = data.get("playbook", {})
    
    # Find MITRE technique if specified
    mitre_technique = None
    if playbook_data.get("mitre_technique_id"):
        mitre_technique = MitreAttackTechnique.objects.filter(
            technique_id=playbook_data["mitre_technique_id"]
        ).first()
    
    # Create the graph
    graph = PlaybookGraph.objects.create(
        title=playbook_data.get("title", "Imported Playbook"),
        organization=organization,
        author=author,
        status="IDEA",  # Always start as IDEA
        
        # Strategy
        mitre_technique=mitre_technique,
        selected_strategy=playbook_data.get("selected_strategy", {}),
        detection_rule=playbook_data.get("detection_rule", ""),
        
        # Context
        goal=playbook_data.get("goal", ""),
        technical_context=playbook_data.get("technical_context", ""),
        blind_spots=playbook_data.get("blind_spots", ""),
        triage_guidance=playbook_data.get("triage_guidance", ""),
        false_positives=playbook_data.get("false_positives", ""),
        response_playbook=playbook_data.get("response_playbook", ""),
        
        # Testing
        test_scenario=playbook_data.get("test_scenario", ""),
        test_expected_output=playbook_data.get("test_expected_output", ""),
        
        # Deployment
        target_file_path=playbook_data.get("target_file_path", ""),
        
        # Valuation
        robustness_level=playbook_data.get("robustness_level", 0),
        data_source_robustness=playbook_data.get("data_source_robustness", ""),
        
        # SOAR
        alert_trigger=playbook_data.get("alert_trigger", ""),
        default_severity=playbook_data.get("default_severity", "MEDIUM"),
        enrichment_steps=playbook_data.get("enrichment_steps", []),
        containment_steps=playbook_data.get("containment_steps", []),
        notification_steps=playbook_data.get("notification_steps", []),
        downstream_correlation_requirements=playbook_data.get("downstream_correlation_requirements", {}),
        
        # Notes
        notes=playbook_data.get("notes", "")
    )
    
    # Add tags
    tags = playbook_data.get("tags", [])
    if tags:
        graph.tags.add(*tags)
    
    # Create node ID mapping (old ID -> new node)
    node_id_map = {}
    
    # Create nodes
    for node_data in playbook_data.get("nodes", []):
        old_id = node_data.get("id")
        node = PlaybookNode.objects.create(
            graph=graph,
            layer_name=node_data.get("layer_name", "Node"),
            position_x=node_data.get("position_x", 0),
            position_y=node_data.get("position_y", 0),
            color=node_data.get("color", "default"),
            ui_metadata=node_data.get("ui_metadata", {})
        )
        
        # Add MITRE mappings to node
        mitre_ids = node_data.get("mitre_attack_ids", [])
        if mitre_ids:
            techniques = MitreAttackTechnique.objects.filter(technique_id__in=mitre_ids)
            node.mitre_attack_mappings.set(techniques)
        
        node_id_map[old_id] = node
    
    # Create edges (using the ID mapping)
    for edge_data in playbook_data.get("edges", []):
        source_old_id = edge_data.get("source_node_id")
        target_old_id = edge_data.get("target_node_id")
        
        source_node = node_id_map.get(source_old_id)
        target_node = node_id_map.get(target_old_id)
        
        if source_node and target_node:
            PlaybookEdge.objects.create(
                graph=graph,
                source_node=source_node,
                target_node=target_node
            )
    
    # Log activity
    ActivityLog.objects.create(
        playbook=graph,
        user=author,
        action="IMPORTED",
        details=f"Playbook imported from external file"
    )
    
    return graph


def _clear_graph_nodes_and_edges(graph: PlaybookGraph) -> None:
    graph.edges.all().delete()
    graph.nodes.all().delete()


def _apply_nodes_edges_from_v1(graph: PlaybookGraph, playbook_data: dict) -> None:
    node_id_map = {}

    for node_data in playbook_data.get("nodes", []):
        old_id = node_data.get("id")
        node = PlaybookNode.objects.create(
            graph=graph,
            layer_name=node_data.get("layer_name", "Node"),
            position_x=node_data.get("position_x", 0),
            position_y=node_data.get("position_y", 0),
            color=node_data.get("color", "default"),
            ui_metadata=node_data.get("ui_metadata", {})
        )

        mitre_ids = node_data.get("mitre_attack_ids", [])
        if mitre_ids:
            techniques = MitreAttackTechnique.objects.filter(technique_id__in=mitre_ids)
            node.mitre_attack_mappings.set(techniques)

        node_id_map[old_id] = node

    for edge_data in playbook_data.get("edges", []):
        source_old_id = edge_data.get("source_node_id")
        target_old_id = edge_data.get("target_node_id")

        source_node = node_id_map.get(source_old_id)
        target_node = node_id_map.get(target_old_id)

        if source_node and target_node:
            PlaybookEdge.objects.create(
                graph=graph,
                source_node=source_node,
                target_node=target_node
            )


def update_playbook_graph_from_v1(data: dict, graph: PlaybookGraph, author, new_title: str | None = None) -> PlaybookGraph:
    playbook_data = data.get("playbook", {})

    if new_title:
        playbook_data["title"] = new_title

    mitre_technique = None
    if playbook_data.get("mitre_technique_id"):
        mitre_technique = MitreAttackTechnique.objects.filter(
            technique_id=playbook_data["mitre_technique_id"]
        ).first()

    graph.title = playbook_data.get("title", graph.title)
    graph.mitre_technique = mitre_technique
    graph.selected_strategy = playbook_data.get("selected_strategy", {})
    graph.detection_rule = playbook_data.get("detection_rule", "")

    graph.goal = playbook_data.get("goal", "")
    graph.technical_context = playbook_data.get("technical_context", "")
    graph.blind_spots = playbook_data.get("blind_spots", "")
    graph.triage_guidance = playbook_data.get("triage_guidance", "")
    graph.false_positives = playbook_data.get("false_positives", "")
    graph.response_playbook = playbook_data.get("response_playbook", "")

    graph.test_scenario = playbook_data.get("test_scenario", "")
    graph.test_expected_output = playbook_data.get("test_expected_output", "")
    graph.target_file_path = playbook_data.get("target_file_path", "")

    graph.robustness_level = playbook_data.get("robustness_level", 0)
    graph.data_source_robustness = playbook_data.get("data_source_robustness", "")

    graph.alert_trigger = playbook_data.get("alert_trigger", "")
    graph.default_severity = playbook_data.get("default_severity", "MEDIUM")
    graph.enrichment_steps = playbook_data.get("enrichment_steps", [])
    graph.containment_steps = playbook_data.get("containment_steps", [])
    graph.notification_steps = playbook_data.get("notification_steps", [])
    graph.downstream_correlation_requirements = playbook_data.get("downstream_correlation_requirements", {})

    graph.notes = playbook_data.get("notes", "")
    graph.save()

    graph.tags.set(playbook_data.get("tags", []))

    _clear_graph_nodes_and_edges(graph)
    _apply_nodes_edges_from_v1(graph, playbook_data)

    ActivityLog.objects.create(
        playbook=graph,
        user=author,
        action="IMPORTED",
        details="Playbook overwritten from external file"
    )

    return graph


# =============================================================================
# HEX v2.0 FORMAT - New Standardized Export/Import Format
# =============================================================================

def serialize_playbook_graph_hex_v2(graph: PlaybookGraph) -> dict:
    """
    Serialize a PlaybookGraph to HEX v2.0 format.
    HEX = HEFAISTOS Export format
    v2.0 = Human-readable, developer-friendly, includes capability abstraction
    """
    import json
    from datetime import datetime
    
    # Serialize nodes with layer information
    nodes = []
    layer_to_nodes = {}  # Map layers to node IDs for capability abstraction
    
    for node in graph.nodes.all():
        node_data = {
            "id": str(node.id),
            "name": node.layer_name,
            "type": "detection",  # Could be enhanced with node.node_type
            "layer": None,  # Will be populated from graph structure
            "description": "",  # Could come from node.ui_metadata
            "position": {
                "x": float(node.position_x) if node.position_x else 0,
                "y": float(node.position_y) if node.position_y else 0
            },
            "color": node.color or "#95E1D3",
            "mitre_techniques": [t.technique_id for t in node.mitre_attack_mappings.all()]
        }
        nodes.append(node_data)
    
    # Serialize edges
    edges = []
    for edge in graph.edges.all():
        edge_data = {
            "id": str(edge.id),
            "source": str(edge.source_node_id),
            "target": str(edge.target_node_id),
            "label": "trigger"
        }
        edges.append(edge_data)
    
    # Parse selected_strategy to extract capability layers
    layers = []
    try:
        strategy_obj = json.loads(graph.selected_strategy) if isinstance(graph.selected_strategy, str) else (graph.selected_strategy or {})
        if isinstance(strategy_obj, dict) and "layers" in strategy_obj:
            for layer in strategy_obj.get("layers", []):
                layer_data = {
                    "layer_id": layer.get("id", f"layer_{len(layers)+1}"),
                    "layer_name": layer.get("name", "Unnamed Layer"),
                    "capability": layer.get("capability", ""),
                    "description": layer.get("description", ""),
                    "nodes": layer.get("nodes", [])
                }
                layers.append(layer_data)
    except (json.JSONDecodeError, TypeError):
        # If no structured layers, create a simple single layer
        layers = [{
            "layer_id": "layer_1",
            "layer_name": "Processing Layer",
            "capability": "Detection and Response",
            "description": "All processing steps",
            "nodes": [n["id"] for n in nodes]
        }]
    
    # Build HEX v2.0 export object
    export_data = {
        "hex_format": "2.0",
        "metadata": {
            "name": graph.title,
            "description": graph.goal or "",
            "version": str(graph.version or "1.0.0"),
            "status": graph.status or "DEVELOPMENT",
            "tags": list(graph.tags.names()),
            "created_by": graph.author.username if graph.author else "unknown",
            "created_date": graph.created_at.isoformat() if graph.created_at else datetime.now().isoformat(),
            "last_modified": graph.updated_at.isoformat() if hasattr(graph, 'updated_at') and graph.updated_at else datetime.now().isoformat()
        },
        "strategy": {
            "mitre_techniques": [
                {
                    "technique_id": graph.mitre_technique.technique_id,
                    "name": graph.mitre_technique.name,
                    "tactic": graph.mitre_technique.tactic if hasattr(graph.mitre_technique, 'tactic') else ""
                }
            ] if graph.mitre_technique else [],
            "detection_approach": graph.selected_strategy if isinstance(graph.selected_strategy, str) else json.dumps(graph.selected_strategy or {}),
            "selected_detection_method": "Detection Logic Defined Below"
        },
        "capability_abstraction": {
            "mission": {
                "goal": graph.goal or "",
                "description": graph.technical_context or ""
            },
            "layers": layers
        },
        "detection_logic": {
            "detection_rule": graph.detection_rule or "",
            "rule_format": "kql",
            "data_sources": [],
            "blind_spots": [s.strip() for s in (graph.blind_spots or "").split('\n') if s.strip()]
        },
        "operational_context": {
            "goal": graph.goal or "",
            "technical_context": graph.technical_context or "",
            "false_positives": [fp.strip() for fp in (graph.false_positives or "").split('\n') if fp.strip()],
            "triage_guidance": graph.triage_guidance or "",
            "response_playbook": graph.response_playbook or ""
        },
        "testing": {
            "test_scenario": graph.test_scenario or "",
            "test_expected_output": graph.test_expected_output or "",
            "test_environment": "",
            "target_file_path": graph.target_file_path or ""
        },
        "soar_configuration": {
            "alert_trigger": graph.alert_trigger or "",
            "default_severity": graph.default_severity or "MEDIUM",
            "enrichment_steps": graph.enrichment_steps or [],
            "containment_steps": graph.containment_steps or [],
            "notification_steps": graph.notification_steps or [],
            "downstream_correlation_requirements": graph.downstream_correlation_requirements or {},
        },
        "graph_structure": {
            "nodes": nodes,
            "edges": edges
        },
        "audit_trail": {
            "robustness_level": graph.robustness_level or 0,
            "data_source_robustness": graph.data_source_robustness or "",
            "data_source_maturity": "",
            "notes": graph.notes or "",
            "validation_status": "Not validated"
        }
    }
    
    return export_data


def deserialize_playbook_graph_hex_v2(data: dict, organization, author) -> PlaybookGraph:
    """
    Create a new PlaybookGraph from HEX v2.0 format data.
    Returns the created graph.
    """
    import json
    
    metadata = data.get("metadata", {})
    strategy = data.get("strategy", {})
    detection_logic = data.get("detection_logic", {})
    operational_context = data.get("operational_context", {})
    testing = data.get("testing", {})
    soar_config = data.get("soar_configuration", {})
    graph_structure = data.get("graph_structure", {})
    audit_trail = data.get("audit_trail", {})
    capability_abstraction = data.get("capability_abstraction", {})
    
    # Find MITRE technique if specified
    mitre_technique = None
    mitre_techniques = strategy.get("mitre_techniques", [])
    if mitre_techniques and isinstance(mitre_techniques, list) and len(mitre_techniques) > 0:
        technique_id = mitre_techniques[0].get("technique_id")
        if technique_id:
            mitre_technique = MitreAttackTechnique.objects.filter(
                technique_id=technique_id
            ).first()
    
    # Prepare selected_strategy with layers from capability_abstraction
    layers = capability_abstraction.get("layers", [])
    selected_strategy = {
        "layers": [
            {
                "id": layer.get("layer_id"),
                "name": layer.get("layer_name"),
                "capability": layer.get("capability"),
                "description": layer.get("description"),
                "nodes": layer.get("nodes", [])
            }
            for layer in layers
        ]
    }
    
    # Create the graph
    graph = PlaybookGraph.objects.create(
        title=metadata.get("name", "Imported Playbook"),
        organization=organization,
        author=author,
        status=metadata.get("status", "DEVELOPMENT"),
        
        # Strategy
        mitre_technique=mitre_technique,
        selected_strategy=json.dumps(selected_strategy),
        detection_rule=detection_logic.get("detection_rule", ""),
        
        # Operational Context
        goal=operational_context.get("goal", ""),
        technical_context=operational_context.get("technical_context", ""),
        blind_spots="\n".join(detection_logic.get("blind_spots", [])),
        triage_guidance=operational_context.get("triage_guidance", ""),
        false_positives="\n".join(operational_context.get("false_positives", [])),
        response_playbook=operational_context.get("response_playbook", ""),
        
        # Testing
        test_scenario=testing.get("test_scenario", ""),
        test_expected_output=testing.get("test_expected_output", ""),
        
        # Deployment
        target_file_path=testing.get("target_file_path", ""),
        
        # Audit & Valuation
        robustness_level=audit_trail.get("robustness_level", 0),
        data_source_robustness=audit_trail.get("data_source_robustness", ""),
        
        # SOAR Configuration
        alert_trigger=soar_config.get("alert_trigger", ""),
        default_severity=soar_config.get("default_severity", "MEDIUM"),
        enrichment_steps=soar_config.get("enrichment_steps", []),
        containment_steps=soar_config.get("containment_steps", []),
        notification_steps=soar_config.get("notification_steps", []),
        downstream_correlation_requirements=soar_config.get("downstream_correlation_requirements", {}),
        
        # Notes
        notes=audit_trail.get("notes", "")
    )
    
    # Add tags
    tags = metadata.get("tags", [])
    if tags:
        graph.tags.add(*tags)
    
    # Create node ID mapping (old ID -> new node)
    node_id_map = {}
    nodes_data = graph_structure.get("nodes", [])
    
    # Create nodes
    for node_data in nodes_data:
        old_id = node_data.get("id")
        position = node_data.get("position", {})
        node = PlaybookNode.objects.create(
            graph=graph,
            layer_name=node_data.get("name", "Node"),
            position_x=position.get("x", 0),
            position_y=position.get("y", 0),
            color=node_data.get("color", "#95E1D3"),
            ui_metadata={"type": node_data.get("type", "detection")}
        )
        
        # Add MITRE mappings to node
        mitre_ids = node_data.get("mitre_techniques", [])
        if mitre_ids:
            techniques = MitreAttackTechnique.objects.filter(technique_id__in=mitre_ids)
            node.mitre_attack_mappings.set(techniques)
        
        node_id_map[old_id] = node
    
    # Create edges (using the ID mapping)
    edges_data = graph_structure.get("edges", [])
    for edge_data in edges_data:
        source_id = edge_data.get("source")
        target_id = edge_data.get("target")
        
        source_node = node_id_map.get(source_id)
        target_node = node_id_map.get(target_id)
        
        if source_node and target_node:
            PlaybookEdge.objects.create(
                graph=graph,
                source_node=source_node,
                target_node=target_node
            )
    
    # Log activity
    ActivityLog.objects.create(
        playbook=graph,
        user=author,
        action="IMPORTED",
        details=f"Playbook imported from HEX v2.0 format"
    )
    
    return graph


def _apply_nodes_edges_from_hex_v2(graph: PlaybookGraph, graph_structure: dict) -> None:
    node_id_map = {}
    nodes_data = graph_structure.get("nodes", [])

    for node_data in nodes_data:
        old_id = node_data.get("id")
        position = node_data.get("position", {})
        node = PlaybookNode.objects.create(
            graph=graph,
            layer_name=node_data.get("name", "Node"),
            position_x=position.get("x", 0),
            position_y=position.get("y", 0),
            color=node_data.get("color", "#95E1D3"),
            ui_metadata={"type": node_data.get("type", "detection")}
        )

        mitre_ids = node_data.get("mitre_techniques", [])
        if mitre_ids:
            techniques = MitreAttackTechnique.objects.filter(technique_id__in=mitre_ids)
            node.mitre_attack_mappings.set(techniques)

        node_id_map[old_id] = node

    edges_data = graph_structure.get("edges", [])
    for edge_data in edges_data:
        source_id = edge_data.get("source")
        target_id = edge_data.get("target")

        source_node = node_id_map.get(source_id)
        target_node = node_id_map.get(target_id)

        if source_node and target_node:
            PlaybookEdge.objects.create(
                graph=graph,
                source_node=source_node,
                target_node=target_node
            )


def update_playbook_graph_from_hex_v2(data: dict, graph: PlaybookGraph, author, new_title: str | None = None) -> PlaybookGraph:
    import json

    metadata = data.get("metadata", {})
    strategy = data.get("strategy", {})
    detection_logic = data.get("detection_logic", {})
    operational_context = data.get("operational_context", {})
    testing = data.get("testing", {})
    soar_config = data.get("soar_configuration", {})
    graph_structure = data.get("graph_structure", {})
    audit_trail = data.get("audit_trail", {})
    capability_abstraction = data.get("capability_abstraction", {})

    if new_title:
        metadata["name"] = new_title

    mitre_technique = None
    mitre_techniques = strategy.get("mitre_techniques", [])
    if mitre_techniques and isinstance(mitre_techniques, list) and len(mitre_techniques) > 0:
        technique_id = mitre_techniques[0].get("technique_id")
        if technique_id:
            mitre_technique = MitreAttackTechnique.objects.filter(
                technique_id=technique_id
            ).first()

    layers = capability_abstraction.get("layers", [])
    selected_strategy = {
        "layers": [
            {
                "id": layer.get("layer_id"),
                "name": layer.get("layer_name"),
                "capability": layer.get("capability"),
                "description": layer.get("description"),
                "nodes": layer.get("nodes", [])
            }
            for layer in layers
        ]
    }

    graph.title = metadata.get("name", graph.title)
    graph.status = metadata.get("status", graph.status or "DEVELOPMENT")
    graph.mitre_technique = mitre_technique
    graph.selected_strategy = json.dumps(selected_strategy)
    graph.detection_rule = detection_logic.get("detection_rule", "")

    graph.goal = operational_context.get("goal", "")
    graph.technical_context = operational_context.get("technical_context", "")
    graph.blind_spots = "\n".join(detection_logic.get("blind_spots", []))
    graph.triage_guidance = operational_context.get("triage_guidance", "")
    graph.false_positives = "\n".join(operational_context.get("false_positives", []))
    graph.response_playbook = operational_context.get("response_playbook", "")

    graph.test_scenario = testing.get("test_scenario", "")
    graph.test_expected_output = testing.get("test_expected_output", "")
    graph.target_file_path = testing.get("target_file_path", "")

    graph.robustness_level = audit_trail.get("robustness_level", 0)
    graph.data_source_robustness = audit_trail.get("data_source_robustness", "")

    graph.alert_trigger = soar_config.get("alert_trigger", "")
    graph.default_severity = soar_config.get("default_severity", "MEDIUM")
    graph.enrichment_steps = soar_config.get("enrichment_steps", [])
    graph.containment_steps = soar_config.get("containment_steps", [])
    graph.notification_steps = soar_config.get("notification_steps", [])
    graph.downstream_correlation_requirements = soar_config.get("downstream_correlation_requirements", {})

    graph.notes = audit_trail.get("notes", "")
    graph.save()

    graph.tags.set(metadata.get("tags", []))

    _clear_graph_nodes_and_edges(graph)
    _apply_nodes_edges_from_hex_v2(graph, graph_structure)

    ActivityLog.objects.create(
        playbook=graph,
        user=author,
        action="IMPORTED",
        details="Playbook overwritten from HEX v2.0 format"
    )

    return graph


class ExportPlaybookGraph(graphene.Mutation):
    """Export a PlaybookGraph as JSON"""
    
    class Arguments:
        graph_id = graphene.UUID(required=True)
    
    success = graphene.Boolean()
    export_data = graphene.JSONString()
    message = graphene.String()
    
    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id):
        user = info.context.user
        
        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            return ExportPlaybookGraph(success=False, message="Playbook not found", export_data=None)
        
        # Export in HEX v2.0 format (new standardized format)
        export_data = serialize_playbook_graph_hex_v2(graph)
        
        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="EXPORTED",
            details="Playbook exported to HEX v2.0 JSON format"
        )
        
        return ExportPlaybookGraph(
            success=True,
            export_data=json.dumps(export_data, indent=2),
            message="Playbook exported successfully in HEX v2.0 format"
        )


def _build_workbench_document_data(graph):
    """Return a plain dict with the document sections for a PlaybookGraph."""
    return {
        "title": graph.title,
        "id": graph.custom_id or str(graph.id),
        "technical_context": graph.technical_context or "",
        "response_playbook": graph.response_playbook or "",
        "tags": list(graph.tags.names()),
    }


def _export_workbench_docx(data: dict) -> bytes:
    """Generate a DOCX document from workbench data and return raw bytes."""
    from docx import Document
    from io import BytesIO

    doc = Document()

    # Title
    doc.add_heading(data["title"], level=1)

    # ID
    p = doc.add_paragraph()
    p.add_run("ID: ").bold = True
    p.add_run(data["id"])

    # Technical Context
    doc.add_heading("Technical Context", level=2)
    doc.add_paragraph(data["technical_context"] or "N/A")

    # Response Playbook
    doc.add_heading("Response Playbook", level=2)
    doc.add_paragraph(data["response_playbook"] or "N/A")

    # Tags
    doc.add_heading("Tags", level=2)
    doc.add_paragraph(", ".join(data["tags"]) if data["tags"] else "N/A")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_workbench_pdf(data: dict) -> bytes:
    """Generate a PDF document from workbench data and return raw bytes."""
    import html
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from io import BytesIO

    def safe(text: str) -> str:
        """Escape HTML special characters for use in reportlab Paragraph."""
        return html.escape(text or "N/A")

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    heading1 = styles["Heading1"]
    heading2 = styles["Heading2"]
    body = styles["BodyText"]

    story = []

    # Title
    story.append(Paragraph(safe(data["title"]), heading1))
    story.append(Spacer(1, 0.3 * cm))

    # ID
    story.append(Paragraph(f"<b>ID:</b> {safe(data['id'])}", body))
    story.append(Spacer(1, 0.5 * cm))

    # Technical Context
    story.append(Paragraph("Technical Context", heading2))
    story.append(Paragraph(safe(data["technical_context"]), body))
    story.append(Spacer(1, 0.5 * cm))

    # Response Playbook
    story.append(Paragraph("Response Playbook", heading2))
    story.append(Paragraph(safe(data["response_playbook"]), body))
    story.append(Spacer(1, 0.5 * cm))

    # Tags
    story.append(Paragraph("Tags", heading2))
    tags_text = ", ".join(data["tags"]) if data["tags"] else ""
    story.append(Paragraph(safe(tags_text), body))

    doc.build(story)
    return buf.getvalue()


def _export_workbench_csv(data: dict) -> bytes:
    """Generate a CSV document from workbench data and return raw bytes."""
    import csv
    from io import StringIO

    buf = StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Field", "Value"])
    writer.writerow(["Title", data["title"]])
    writer.writerow(["ID", data["id"]])
    writer.writerow(["Technical Context", data["technical_context"]])
    writer.writerow(["Response Playbook", data["response_playbook"]])
    writer.writerow(["Tags", ", ".join(data["tags"])])
    return buf.getvalue().encode("utf-8")


class ExportWorkbenchDocument(graphene.Mutation):
    """Export a PlaybookGraph as a formatted document (DOCX, PDF, or CSV)."""

    class Arguments:
        graph_id = graphene.UUID(required=True)
        format = graphene.String(required=True, description="Export format: docx, pdf, or csv")

    success = graphene.Boolean()
    file_data = graphene.String(description="Base64-encoded file content")
    filename = graphene.String()
    content_type = graphene.String()
    message = graphene.String()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id, format):
        user = info.context.user

        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            return ExportWorkbenchDocument(success=False, message="Playbook not found")

        fmt = format.lower()
        if fmt not in ("docx", "pdf", "csv"):
            return ExportWorkbenchDocument(success=False, message=f"Unsupported format: {format}")

        data = _build_workbench_document_data(graph)
        safe_title = re.sub(r"[^a-zA-Z0-9_-]", "_", graph.title)

        try:
            if fmt == "docx":
                raw = _export_workbench_docx(data)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                filename = f"{safe_title}.docx"
            elif fmt == "pdf":
                raw = _export_workbench_pdf(data)
                content_type = "application/pdf"
                filename = f"{safe_title}.pdf"
            else:
                raw = _export_workbench_csv(data)
                content_type = "text/csv"
                filename = f"{safe_title}.csv"
        except Exception as exc:
            export_logger.exception("Document export failed for graph %s", graph_id)
            return ExportWorkbenchDocument(success=False, message=f"Export failed: {exc}")

        file_data = base64.b64encode(raw).decode("utf-8")

        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="EXPORTED",
            details=f"Workbench exported as {fmt.upper()} document"
        )

        return ExportWorkbenchDocument(
            success=True,
            file_data=file_data,
            filename=filename,
            content_type=content_type,
            message=f"Workbench exported as {fmt.upper()} successfully",
        )


class ExportAllWorkbenchesHexV2(graphene.Mutation):
    """Export every workbench in the caller's organization as a HEX v2.0 ZIP archive."""

    success = graphene.Boolean()
    file_data = graphene.String(description="Base64-encoded zip content")
    filename = graphene.String()
    content_type = graphene.String()
    message = graphene.String()

    @staticmethod
    @role_required([Roles.ADMIN])
    def mutate(root, info):
        from io import BytesIO
        import zipfile

        user = info.context.user
        exported_at = timezone.now()
        graphs = PlaybookGraph.objects.filter(
            organization=user.organization,
        ).select_related(
            'author',
            'mitre_technique',
        ).prefetch_related(
            'tags',
            'edges',
            'nodes__mitre_attack_mappings',
        )

        manifest_entries = []
        archive_buffer = BytesIO()
        with zipfile.ZipFile(archive_buffer, 'w', compression=zipfile.ZIP_DEFLATED) as archive:
            for graph in graphs:
                safe_title = re.sub(r"[^a-zA-Z0-9_-]", "_", graph.title or '') or 'workbench'
                entry_filename = f'{safe_title}__{graph.id}.hex.json'
                archive.writestr(
                    entry_filename,
                    json.dumps(serialize_playbook_graph_hex_v2(graph), indent=2),
                )
                manifest_entries.append({
                    'id': str(graph.id),
                    'title': graph.title,
                    'filename': entry_filename,
                    'version': graph.version,
                    'minor_version': graph.minor_version,
                })

            archive.writestr(
                'manifest.json',
                json.dumps(
                    {
                        'exported_at': exported_at.isoformat(),
                        'organization': getattr(user.organization, 'name', ''),
                        'count': len(manifest_entries),
                        'hex_format': '2.0',
                        'entries': manifest_entries,
                    },
                    indent=2,
                ),
            )

        org_name = re.sub(r"[^a-zA-Z0-9_-]", "_", getattr(user.organization, 'name', '') or 'organization')
        filename = f'workbenches_hex_v2_{org_name}_{exported_at.strftime("%Y%m%dT%H%M%SZ")}.zip'
        return ExportAllWorkbenchesHexV2(
            success=True,
            file_data=base64.b64encode(archive_buffer.getvalue()).decode('utf-8'),
            filename=filename,
            content_type='application/zip',
            message='Workbench archive exported successfully',
        )


class ImportPlaybookGraph(graphene.Mutation):
    """Import a PlaybookGraph from HEX v2.0 JSON format"""
    
    class Arguments:
        import_data = graphene.JSONString(required=True, description="JSON string of exported playbook in HEX v2.0 format")
        new_title = graphene.String(description="Optional: Override the title")
        graph_id = graphene.UUID(description="Optional: Overwrite an existing playbook graph")
    
    success = graphene.Boolean()
    graph = graphene.Field(PlaybookGraphType)
    message = graphene.String()
    
    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, import_data, new_title=None, graph_id=None):
        user = info.context.user

        target_graph = None
        if graph_id:
            try:
                target_graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
            except PlaybookGraph.DoesNotExist:
                return ImportPlaybookGraph(success=False, message="Playbook not found", graph=None)
        
        try:
            data = json.loads(import_data) if isinstance(import_data, str) else import_data
        except json.JSONDecodeError as e:
            return ImportPlaybookGraph(success=False, message=f"Invalid JSON: {e}", graph=None)
        
        # Detect and validate format
        hex_format = data.get("hex_format")
        export_type = data.get("export_type")
        
        if hex_format == "2.0":
            # HEX v2.0 format
            try:
                # Validate required metadata
                if not data.get("metadata", {}).get("name"):
                    return ImportPlaybookGraph(
                        success=False, 
                        message="Invalid HEX v2.0 format: missing 'metadata.name'",
                        graph=None
                    )
                
                # Override title if provided
                if new_title:
                    data["metadata"]["name"] = new_title
                
                if target_graph:
                    graph = update_playbook_graph_from_hex_v2(data, target_graph, user, new_title)
                else:
                    graph = deserialize_playbook_graph_hex_v2(data, user.organization, user)
                return ImportPlaybookGraph(
                    success=True,
                    graph=graph,
                    message=f"Playbook '{graph.title}' imported successfully from HEX v2.0 format"
                )
            except Exception as e:
                return ImportPlaybookGraph(
                    success=False,
                    message=f"Error importing HEX v2.0 playbook: {str(e)}",
                    graph=None
                )
        
        elif export_type == "playbook_graph":
            # Legacy V1 format - still supported
            if new_title:
                data.setdefault("playbook", {})["title"] = new_title
            
            try:
                if target_graph:
                    graph = update_playbook_graph_from_v1(data, target_graph, user, new_title)
                else:
                    graph = deserialize_playbook_graph(data, user.organization, user)
                return ImportPlaybookGraph(
                    success=True,
                    graph=graph,
                    message=f"Playbook '{graph.title}' imported successfully (V1 legacy format)"
                )
            except Exception as e:
                return ImportPlaybookGraph(
                    success=False,
                    message=f"Import failed: {str(e)}",
                    graph=None
                )
        
        else:
            return ImportPlaybookGraph(
                success=False,
                message="Unknown format. Expected 'hex_format: 2.0' or 'export_type: playbook_graph'",
                graph=None
            )


def deserialize_playbook_graph_from_opentide(
    tvm_data: dict,
    dom_data: dict,
    mdr_data: dict,
    organization,
    author,
    new_title: str = None,  # type: ignore[assignment]
) -> 'PlaybookGraph':
    """
    Create a new PlaybookGraph by translating OpenTide TVM / DOM / MDR YAML data
    back into HEFAISTOS workbench fields.

    Field mapping summary
    ─────────────────────
    TVM → title (description), mitre_technique, technical_context, blind_spots
    DOM → goal (description), false_positives, triage_guidance, alert_trigger,
           response_playbook, default_severity (priority), robustness_level,
           data_source_maturity
    MDR → goal (description fallback), default_severity (alert_severity),
           robustness_level (testing.robustness_level),
           DetectionRule objects for each configuration platform

    Args:
        tvm_data:     Parsed TVM YAML dict (may be empty).
        dom_data:     Parsed DOM YAML dict (may be empty).
        mdr_data:     Parsed MDR YAML dict (may be empty).
        organization: Organization instance for the new workbench.
        author:       CustomUser instance who is performing the import.
        new_title:    Optional title override.

    Returns:
        PlaybookGraph: Newly created workbench.
    """
    tvm = tvm_data or {}
    dom = dom_data or {}
    mdr = mdr_data or {}

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    title = new_title
    if not title:
        # TVM description is the most human-readable representation of the title
        title = (
            tvm.get('description')
            or mdr.get('description')
            or dom.get('description')
        )
    if not title:
        # Fall back: derive from name slug, stripping object-type prefix
        slug = mdr.get('name') or dom.get('name') or tvm.get('name') or ''
        parts = slug.split('_')
        if parts and parts[0].lower() in ('mdr', 'dom', 'tvm', 'bdr'):
            parts = parts[1:]
        # Short alphanumeric tokens (up to 6 chars, e.g. "t1070", "de") are
        # rendered in uppercase; longer parts are title-cased as prose words.
        _SHORT_TOKEN_RE = re.compile(r'^[a-z0-9]{1,6}$')
        title = ' '.join(
            p.upper() if _SHORT_TOKEN_RE.match(p) else p.capitalize()
            for p in parts
        ) if parts else 'Imported from OpenTide'
    title = (title or 'Imported from OpenTide').strip()[:255]

    # ------------------------------------------------------------------
    # MITRE technique
    # ------------------------------------------------------------------
    mitre_technique = None
    tvm_mitre = tvm.get('mitre') or {}
    technique_id = tvm_mitre.get('technique_id')
    if technique_id:
        mitre_technique = MitreAttackTechnique.objects.filter(
            technique_id=technique_id
        ).first()

    # ------------------------------------------------------------------
    # Textual workbench fields
    # ------------------------------------------------------------------
    goal = dom.get('description') or mdr.get('description') or ''
    technical_context = tvm.get('technical_context', '')
    blind_spots = tvm.get('blind_spots', '')
    triage_guidance = dom.get('triage_guidance', '') or tvm.get('triage_guidance', '')
    false_positives = dom.get('false_positives', '')

    dom_response = dom.get('response') or {}
    response_playbook = dom_response.get('playbook', '')
    alert_trigger = dom_response.get('alert_trigger', '')

    # ------------------------------------------------------------------
    # Severity
    # Priority labels used in DOM → internal severity codes
    # ------------------------------------------------------------------
    _PRIORITY_TO_SEVERITY = {
        'Critical incident': 'CRITICAL',
        'Critical': 'CRITICAL',
        'High-severity incident': 'HIGH',
        'High': 'HIGH',
        'Moderate incident': 'MEDIUM',
        'Medium': 'MEDIUM',
        'Low-severity incident': 'LOW',
        'Low': 'LOW',
        'Informational event': 'INFORMATIONAL',
        'Informational': 'INFORMATIONAL',
    }
    mdr_response = mdr.get('response') or {}
    raw_severity = (
        mdr_response.get('alert_severity')
        or _PRIORITY_TO_SEVERITY.get(dom.get('priority', ''))
        or dom_response.get('severity', '')
        or 'MEDIUM'
    )
    default_severity = (raw_severity or 'MEDIUM').upper()
    _VALID_SEVERITIES = {'CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL'}
    if default_severity not in _VALID_SEVERITIES:
        default_severity = 'MEDIUM'

    # ------------------------------------------------------------------
    # Robustness level & data source maturity
    # ------------------------------------------------------------------
    mdr_testing = mdr.get('testing') or {}
    dom_validation = dom.get('validation') or {}
    try:
        robustness_level = int(
            mdr_testing.get('robustness_level')
            or dom_validation.get('robustness_level')
            or 0
        )
        if robustness_level < 0 or robustness_level > 5:
            robustness_level = 0
    except (TypeError, ValueError):
        robustness_level = 0

    data_source_maturity = dom_validation.get('data_source_maturity', '')

    # ------------------------------------------------------------------
    # TVM v2.1 fields: TLP, references, threat surface, threat actors
    # ------------------------------------------------------------------
    tvm_metadata = tvm.get('metadata') or {}
    tlp_classification = tvm_metadata.get('tlp', 'AMBER')
    # Normalise to our stored choices (strip "TLP:" prefix if present)
    if tlp_classification.startswith('TLP:'):
        tlp_classification = tlp_classification[4:]
    _VALID_TLPS = {'CLEAR', 'GREEN', 'AMBER', 'AMBER+STRICT', 'RED'}
    if tlp_classification not in _VALID_TLPS:
        tlp_classification = 'AMBER'

    tvm_references = tvm.get('references') or {}
    public_references = list((tvm_references.get('public') or {}).values())
    internal_references = list((tvm_references.get('internal') or {}).values())

    tvm_threat = tvm.get('threat') or {}
    threat_surface = tvm_threat.get('surface') or []
    threat_actors = tvm_threat.get('actors') or []

    # ------------------------------------------------------------------
    # Create the PlaybookGraph
    # ------------------------------------------------------------------
    graph = PlaybookGraph.objects.create(
        title=title,
        organization=organization,
        author=author,
        status='DEVELOPMENT',
        mitre_technique=mitre_technique,
        goal=goal,
        technical_context=technical_context,
        blind_spots=blind_spots,
        triage_guidance=triage_guidance,
        false_positives=false_positives,
        response_playbook=response_playbook,
        alert_trigger=alert_trigger,
        default_severity=default_severity,
        robustness_level=robustness_level,
        data_source_maturity=data_source_maturity,
        tlp_classification=tlp_classification,
        public_references=public_references,
        internal_references=internal_references,
        threat_surface=threat_surface,
        threat_actors=threat_actors,
    )

    # ------------------------------------------------------------------
    # Detection Rules – one per platform configuration in the MDR
    # ------------------------------------------------------------------
    # Maps MDR configuration key → (HEFAISTOS format, YAML content key)
    _CONFIG_TO_FORMAT = {
        'defender_for_endpoint': ('KQL', 'query'),
        'splunk': ('SPL', 'query'),
        'wazuh': ('WAZUH', 'rule'),
    }

    configurations = mdr.get('configurations') or {}
    if configurations:
        repo = RuleRepository.objects.filter(
            organization=organization,
            name='Rule Repo',
        ).first()
        if repo is None:
            repo = RuleRepository.objects.create(
                organization=organization,
                name='Rule Repo',
                git_url=None,
            )

        for config_key, (fmt, content_key) in _CONFIG_TO_FORMAT.items():
            config_block = configurations.get(config_key) or {}
            raw_content = (config_block.get(content_key) or '').strip()
            if raw_content:
                DetectionRule.objects.create(
                    organization=organization,
                    repository=repo,
                    playbook=graph,
                    title=f"{title}-{fmt.lower()}",
                    format=fmt,
                    raw_content=raw_content,
                    description=goal or title,
                    author=author.username,
                )

    return graph


class ImportFromOpenTide(graphene.Mutation):
    """Import a PlaybookGraph from OpenTide TVM / DOM / MDR YAML content."""

    class Arguments:
        mdr_yaml = graphene.String(
            required=True,
            description="MDR (Managed Detection Rule) YAML content – required",
        )
        tvm_yaml = graphene.String(
            description="Optional TVM (Threat Vector Model) YAML content",
        )
        dom_yaml = graphene.String(
            description="Optional DOM (Detection Objective Model) YAML content",
        )
        new_title = graphene.String(
            description="Optional title override for the imported workbench",
        )

    success = graphene.Boolean()
    graph = graphene.Field(lambda: PlaybookGraphType)
    message = graphene.String()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, mdr_yaml, tvm_yaml=None, dom_yaml=None, new_title=None):
        user = info.context.user

        # Parse MDR YAML (required)
        try:
            mdr_data = yaml.safe_load(mdr_yaml) if mdr_yaml else {}
            if not isinstance(mdr_data, dict):
                return ImportFromOpenTide(
                    success=False, message="MDR YAML must be a YAML mapping", graph=None
                )
        except yaml.YAMLError as exc:
            return ImportFromOpenTide(
                success=False, message=f"Invalid MDR YAML: {exc}", graph=None
            )

        # Parse TVM YAML (optional)
        tvm_data = None
        if tvm_yaml:
            try:
                tvm_data = yaml.safe_load(tvm_yaml)
                if not isinstance(tvm_data, dict):
                    tvm_data = None
            except yaml.YAMLError as exc:
                return ImportFromOpenTide(
                    success=False, message=f"Invalid TVM YAML: {exc}", graph=None
                )

        # Parse DOM YAML (optional)
        dom_data = None
        if dom_yaml:
            try:
                dom_data = yaml.safe_load(dom_yaml)
                if not isinstance(dom_data, dict):
                    dom_data = None
            except yaml.YAMLError as exc:
                return ImportFromOpenTide(
                    success=False, message=f"Invalid DOM YAML: {exc}", graph=None
                )

        if not mdr_data and not tvm_data and not dom_data:
            return ImportFromOpenTide(
                success=False, message="No valid OpenTide YAML data provided", graph=None
            )

        try:
            graph = deserialize_playbook_graph_from_opentide(
                tvm_data=tvm_data or {},
                dom_data=dom_data or {},
                mdr_data=mdr_data or {},
                organization=user.organization,
                author=user,
                new_title=new_title,
            )

            ActivityLog.objects.create(
                playbook=graph,
                user=user,
                action="IMPORTED_FROM_OPENTIDE",
                details="Workbench imported from OpenTide YAML",
            )

            return ImportFromOpenTide(
                success=True,
                graph=graph,
                message=f"Workbench '{graph.title}' imported successfully from OpenTide",
            )
        except Exception as exc:
            export_logger.error("OpenTide import failed: %s", exc)
            return ImportFromOpenTide(
                success=False, message=f"Import failed: {exc}", graph=None
            )


class PushPlaybookToGitHub(graphene.Mutation):
    """Push a PlaybookGraph to a GitHub repository"""
    
    class Arguments:
        graph_id = graphene.UUID(required=True)
        repository_id = graphene.String(description="ID of a configured RuleRepository to use (from /mgmt/config)")
        github_token = graphene.String(description="GitHub personal access token (required if no repositoryId)")
        repo_owner = graphene.String(description="GitHub repo owner (required if no repositoryId)")
        repo_name = graphene.String(description="GitHub repo name (required if no repositoryId)")
        file_path = graphene.String(description="Path in repo (default: playbooks/{title}.json)")
        branch = graphene.String(description="Branch name (default: main)")
        commit_message = graphene.String(description="Commit message")
        push_opentide_bundle = graphene.Boolean(
            default_value=True,
            description="Push the full OpenTide HEF Objects bundle (TVM/DOM/MDR). Default: true."
        )
        push_platform_rules = graphene.Boolean(
            default_value=False,
            description="Also extract and push individual platform rule files into format directories (kql/, splunk/, sigma/, etc.). Default: false."
        )
    
    success = graphene.Boolean()
    message = graphene.String()
    url = graphene.String(description="URL to the created/updated file on GitHub")

    @staticmethod
    def _extract_github_details(repo_url):
        if not repo_url:
            return None

        patterns = [
            r'^https?://github\.com/([^/]+)/([^/]+?)(?:\.git)?/?$',
            r'^git@github\.com:([^/]+)/([^/]+?)(?:\.git)?$',
        ]
        for pattern in patterns:
            match = re.match(pattern, repo_url.strip())
            if match:
                return match.group(1), match.group(2)
        return None

    @staticmethod
    def _join_repo_path(*parts):
        clean_parts = [part.strip('/') for part in parts if part and str(part).strip('/')]
        return posixpath.join(*clean_parts) if clean_parts else ''

    @staticmethod
    def _build_opentide_bundle(graph, target_folder=None):
        from playbooks.git_client import sanitize_filename
        from playbooks.utils.opentide_compiler import (
            compile_dom_yaml,
            compile_mdr_yaml,
            compile_tvm_yaml,
            dump_opentide_yaml,
            _normalize_mdr_impacted_entities,
        )
        from playbooks.utils.opentide_validator import (
            validate_dom_structure,
            validate_mdr_structure,
            validate_tvm_structure,
        )

        tvm_data = compile_tvm_yaml(graph)
        dom_data = compile_dom_yaml(graph)
        mdr_data = compile_mdr_yaml(graph)
        _normalize_mdr_impacted_entities(mdr_data)

        validations = [
            ('TVM',) + validate_tvm_structure(tvm_data),
            ('DOM',) + validate_dom_structure(dom_data),
            ('MDR',) + validate_mdr_structure(mdr_data),
        ]
        validation_errors = []
        for label, is_valid, errors in validations:
            if not is_valid:
                validation_errors.extend([f'{label}: {error}' for error in errors])

        if validation_errors:
            return None, validation_errors

        base_folder = (target_folder or '').strip('/')
        files = {
            PushPlaybookToGitHub._join_repo_path(
                base_folder,
                'Objects/Threat Vectors',
                f"{sanitize_filename(tvm_data['name'])}.yaml",
            ): dump_opentide_yaml(tvm_data),
            PushPlaybookToGitHub._join_repo_path(
                base_folder,
                'Objects/Detection Objectives',
                f"{sanitize_filename(dom_data['name'])}.yaml",
            ): dump_opentide_yaml(dom_data),
            PushPlaybookToGitHub._join_repo_path(
                base_folder,
                'Objects/Detection Rules',
                f"{sanitize_filename(mdr_data['name'])}.yaml",
            ): dump_opentide_yaml(mdr_data),
        }

        return {
            'files': files,
            'primary_path': next(path for path in files if 'Objects/Detection Rules/' in path),
        }, []

    @staticmethod
    def _build_platform_rule_files(graph, target_folder=None):
        from playbooks.hef_publish import compile_platform_rule_files

        return compile_platform_rule_files(graph, target_folder)

    @staticmethod
    def _create_github_commit(repo_owner, repo_name, branch, github_token, files, commit_message):
        headers = {
            'Authorization': f'token {github_token}',
            'Accept': 'application/vnd.github+json',
            'X-GitHub-Api-Version': '2022-11-28',
        }

        ref_url = f'https://api.github.com/repos/{repo_owner}/{repo_name}/git/ref/heads/{branch}'
        ref_resp = requests.get(ref_url, headers=headers)
        if ref_resp.status_code != 200:
            error_msg = ref_resp.json().get('message', ref_resp.text)
            raise ValueError(f'Unable to access branch {branch}: {error_msg}')

        base_commit_sha = ref_resp.json().get('object', {}).get('sha')
        if not base_commit_sha:
            raise ValueError(f'Unable to resolve HEAD SHA for branch {branch}')

        commit_url = f'https://api.github.com/repos/{repo_owner}/{repo_name}/git/commits/{base_commit_sha}'
        commit_resp = requests.get(commit_url, headers=headers)
        if commit_resp.status_code != 200:
            error_msg = commit_resp.json().get('message', commit_resp.text)
            raise ValueError(f'Unable to read base commit: {error_msg}')

        base_tree_sha = commit_resp.json().get('tree', {}).get('sha')
        if not base_tree_sha:
            raise ValueError('Unable to resolve base tree SHA for repository branch')

        tree_entries = []
        for path, content in files.items():
            blob_resp = requests.post(
                f'https://api.github.com/repos/{repo_owner}/{repo_name}/git/blobs',
                headers=headers,
                json={'content': content, 'encoding': 'utf-8'},
            )
            if blob_resp.status_code not in (200, 201):
                error_msg = blob_resp.json().get('message', blob_resp.text)
                raise ValueError(f'Unable to create blob for {path}: {error_msg}')

            blob_sha = blob_resp.json().get('sha')
            tree_entries.append({
                'path': path,
                'mode': '100644',
                'type': 'blob',
                'sha': blob_sha,
            })

        tree_resp = requests.post(
            f'https://api.github.com/repos/{repo_owner}/{repo_name}/git/trees',
            headers=headers,
            json={'base_tree': base_tree_sha, 'tree': tree_entries},
        )
        if tree_resp.status_code not in (200, 201):
            error_msg = tree_resp.json().get('message', tree_resp.text)
            raise ValueError(f'Unable to create git tree: {error_msg}')

        new_tree_sha = tree_resp.json().get('sha')

        new_commit_resp = requests.post(
            f'https://api.github.com/repos/{repo_owner}/{repo_name}/git/commits',
            headers=headers,
            json={
                'message': commit_message,
                'tree': new_tree_sha,
                'parents': [base_commit_sha],
            },
        )
        if new_commit_resp.status_code not in (200, 201):
            error_msg = new_commit_resp.json().get('message', new_commit_resp.text)
            raise ValueError(f'Unable to create commit: {error_msg}')

        new_commit_sha = new_commit_resp.json().get('sha')

        update_ref_resp = requests.patch(
            ref_url,
            headers=headers,
            json={'sha': new_commit_sha, 'force': False},
        )
        if update_ref_resp.status_code != 200:
            error_msg = update_ref_resp.json().get('message', update_ref_resp.text)
            raise ValueError(f'Unable to update branch ref: {error_msg}')

        return new_commit_sha
    
    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id, repository_id=None, github_token=None, repo_owner=None, repo_name=None, 
               file_path=None, branch="main", commit_message=None,
               push_opentide_bundle=True, push_platform_rules=False):
        user = info.context.user
        repo_url = None
        provider = 'GITHUB'
        api_base_url = None
        
        # Resolve credentials from configured repository if repository_id is provided
        if repository_id:
            try:
                repo = RuleRepository.objects.get(pk=repository_id, organization=user.organization)
            except RuleRepository.DoesNotExist:
                return PushPlaybookToGitHub(success=False, message="Configured repository not found", url=None)
            if not repo.git_url:
                return PushPlaybookToGitHub(success=False, message="Repository has no git URL configured", url=None)
            
            github_token = repo.token
            if not github_token:
                return PushPlaybookToGitHub(success=False, message="Repository has no access token configured", url=None)
            repo_url = repo.git_url
            provider = repo.provider
            api_base_url = repo.api_base_url
        else:
            if not github_token or not repo_owner or not repo_name:
                return PushPlaybookToGitHub(success=False, message="GitHub token, owner, and repo name are required", url=None)
            repo_url = f"https://github.com/{repo_owner}/{repo_name}"
        
        try:
            graph = PlaybookGraph.objects.select_related('organization', 'author', 'mitre_technique').prefetch_related('tags', 'linked_rules').get(
                pk=graph_id,
                organization=user.organization,
            )
        except PlaybookGraph.DoesNotExist:
            return PushPlaybookToGitHub(success=False, message="Playbook not found", url=None)

        all_files = {}
        primary_path = None
        pushed_types = []

        if push_opentide_bundle:
            bundle, validation_errors = PushPlaybookToGitHub._build_opentide_bundle(graph, file_path)
            if validation_errors:
                return PushPlaybookToGitHub(
                    success=False,
                    message='OpenTIDE validation failed: ' + '; '.join(validation_errors),
                    url=None,
                )
            if bundle:
                all_files.update(bundle['files'])
                primary_path = primary_path or bundle.get('primary_path')
                pushed_types.append('bundle')

        if push_platform_rules:
            rule_bundle, rule_errors = PushPlaybookToGitHub._build_platform_rule_files(graph, file_path)
            if rule_errors and not push_opentide_bundle:
                return PushPlaybookToGitHub(
                    success=False,
                    message='Platform rules extraction failed: ' + '; '.join(rule_errors),
                    url=None,
                )
            if rule_bundle:
                all_files.update(rule_bundle['files'])
                primary_path = primary_path or rule_bundle.get('primary_path')
                pushed_types.append('platform_rules')

        if not all_files:
            return PushPlaybookToGitHub(success=False, message='Nothing to push.', url=None)
        
        # Default commit message
        if not commit_message:
            commit_message = f"Publish OpenTIDE HEF package: {graph.title}"
        
        try:
            from playbooks.hef_publish import create_repository_commit

            _commit_sha, client = create_repository_commit(
                repo_url=repo_url,
                token=github_token,
                branch=branch,
                files=all_files,
                commit_message=commit_message,
                provider=provider,
                api_base_url=api_base_url,
            )

            file_url = client.file_web_url(branch, primary_path)

            ActivityLog.objects.create(
                playbook=graph,
                user=user,
                action="PUBLISHED_OPENTIDE_HEF" if push_opentide_bundle else "PUBLISHED_PLATFORM_RULES",
                details=(
                    f"Published {len(all_files)} file(s) [{', '.join(pushed_types)}] "
                    f"to {repo_owner}/{repo_name} ({primary_path or 'repository root'})"
                ),
            )

            return PushPlaybookToGitHub(
                success=True,
                message=f"Published to repository ({len(all_files)} file(s): {', '.join(pushed_types)})",
                url=file_url
            )
        except (requests.RequestException, ValueError) as e:
            return PushPlaybookToGitHub(success=False, message=f"Repository publish failed: {e}", url=None)


class PullPlaybookFromGitHub(graphene.Mutation):
    """Pull (import) a PlaybookGraph from a GitHub repository"""
    
    class Arguments:
        repository_id = graphene.String(description="ID of a configured RuleRepository to use (from /mgmt/config)")
        github_token = graphene.String(description="GitHub personal access token (required if no repositoryId)")
        repo_owner = graphene.String(description="GitHub repo owner (required if no repositoryId)")
        repo_name = graphene.String(description="GitHub repo name (required if no repositoryId)")
        file_path = graphene.String(required=True, description="Path to the JSON file in repo")
        branch = graphene.String(description="Branch name (default: main)")
        new_title = graphene.String(description="Optional: Override the imported title")
        graph_id = graphene.UUID(description="Optional: Overwrite an existing playbook graph")
    
    success = graphene.Boolean()
    graph = graphene.Field(PlaybookGraphType)
    message = graphene.String()
    
    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, file_path, repository_id=None, github_token=None, repo_owner=None, repo_name=None,
               branch="main", new_title=None, graph_id=None):
        user = info.context.user
        repo_url = None
        provider = 'GITHUB'
        api_base_url = None

        # Resolve credentials from configured repository if repository_id is provided
        if repository_id:
            try:
                repo = RuleRepository.objects.get(pk=repository_id, organization=user.organization)
            except RuleRepository.DoesNotExist:
                return PullPlaybookFromGitHub(success=False, message="Configured repository not found", graph=None)
            if not repo.git_url:
                return PullPlaybookFromGitHub(success=False, message="Repository has no git URL configured", graph=None)
            
            github_token = repo.token
            if not github_token:
                return PullPlaybookFromGitHub(success=False, message="Repository has no access token configured", graph=None)
            repo_url = repo.git_url
            provider = repo.provider
            api_base_url = repo.api_base_url
        
        if not repository_id and (not github_token or not repo_owner or not repo_name):
            return PullPlaybookFromGitHub(success=False, message="GitHub token, owner, and repo name are required", graph=None)
        if not repo_url:
            repo_url = f"https://github.com/{repo_owner}/{repo_name}"

        target_graph = None
        if graph_id:
            try:
                target_graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
            except PlaybookGraph.DoesNotExist:
                return PullPlaybookFromGitHub(success=False, message="Playbook not found", graph=None)
        
        try:
            from playbooks.repo_clients import RepoClient

            client = RepoClient(
                repo_url=repo_url,
                token=github_token,
                provider=provider,
                api_base_url=api_base_url,
            )
            ref = client.resolve_commit_sha(branch=branch)
            content = client.get_file_content(file_path, ref)
            if content is None:
                return PullPlaybookFromGitHub(success=False, message="Repository API error: file not found", graph=None)
            
            # Parse JSON
            data = json.loads(content)
            
            hex_format = data.get("hex_format")
            export_type = data.get("export_type")

            if hex_format == "2.0":
                if new_title:
                    data.setdefault("metadata", {})["name"] = new_title

                if target_graph:
                    graph = update_playbook_graph_from_hex_v2(data, target_graph, user, new_title)
                else:
                    graph = deserialize_playbook_graph_hex_v2(data, user.organization, user)
            elif export_type == "playbook_graph":
                if new_title:
                    data.setdefault("playbook", {})["title"] = new_title

                if target_graph:
                    graph = update_playbook_graph_from_v1(data, target_graph, user, new_title)
                else:
                    graph = deserialize_playbook_graph(data, user.organization, user)
            else:
                return PullPlaybookFromGitHub(
                    success=False,
                    message="Invalid file format. Expected HEX v2.0 or legacy playbook export.",
                    graph=None
                )
            
            ActivityLog.objects.create(
                playbook=graph,
                user=user,
                action="PULLED_FROM_GITHUB",
                details=f"Imported from {client.full_name}/{file_path}"
            )
            
            return PullPlaybookFromGitHub(
                success=True,
                graph=graph,
                message=f"Playbook '{graph.title}' imported from repository"
            )
            
        except json.JSONDecodeError as e:
            return PullPlaybookFromGitHub(success=False, message=f"Invalid JSON in file: {e}", graph=None)
        except requests.RequestException as e:
            return PullPlaybookFromGitHub(success=False, message=f"Network error: {e}", graph=None)
        except Exception as e:
            export_logger.error(f"GitHub pull failed: {e}")
            return PullPlaybookFromGitHub(success=False, message=f"Import failed: {e}", graph=None)


class UpdatePlaybookOpenTideYaml(graphene.Mutation):
    class Arguments:
        graph_id = graphene.UUID(required=True)
        opentide_yaml = graphene.JSONString(required=True)
        configured_platforms = graphene.List(graphene.String)

    success = graphene.Boolean()
    playbook_graph = graphene.Field(PlaybookGraphType)

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, graph_id, opentide_yaml, configured_platforms=None):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            graph = PlaybookGraph.objects.get(pk=graph_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            raise Exception("Workbench not found")

        if graph.author != user and not user.is_staff:
            is_admin = getattr(user, 'role', '') in ('ADMIN', 'SUPERADMIN')
            if not is_admin:
                raise Exception("Not authorized to update this workbench")

        import json as _json
        if isinstance(opentide_yaml, str):
            try:
                opentide_data = _json.loads(opentide_yaml)
            except _json.JSONDecodeError:
                opentide_data = opentide_yaml
        else:
            opentide_data = opentide_yaml

        graph.opentide_yaml = opentide_data
        graph.configured_platforms = configured_platforms or []
        graph.save(update_fields=['opentide_yaml', 'configured_platforms', 'updated_at'])

        ActivityLog.objects.create(
            playbook=graph,
            user=user,
            action="OpenTide YAML updated",
            details=f"Configured platforms: {', '.join(graph.configured_platforms) or 'none'}"
        )

        return UpdatePlaybookOpenTideYaml(success=True, playbook_graph=graph)


class RefreshOpenTideMetadata(graphene.Mutation):
    """Manually refresh OpenTide metadata from workbench fields."""

    class Arguments:
        playbook_id = graphene.UUID(required=True)

    success = graphene.Boolean()
    message = graphene.String()
    metadata = graphene.JSONString()

    @staticmethod
    @role_required([Roles.ADMIN, Roles.ANALYST])
    def mutate(root, info, playbook_id):
        user = info.context.user
        if user.is_anonymous:
            raise Exception("Authentication required")

        try:
            playbook = PlaybookGraph.objects.get(pk=playbook_id, organization=user.organization)
        except PlaybookGraph.DoesNotExist:
            return RefreshOpenTideMetadata(success=False, message="Workbench not found", metadata=None)

        if playbook.author != user and not user.is_staff:
            is_admin = getattr(user, 'role', '') in ('ADMIN', 'SUPERADMIN')
            if not is_admin:
                raise Exception("Not authorized to update this workbench")

        playbook.auto_update_opentide_yaml()
        playbook.save(update_fields=['opentide_yaml', 'configured_platforms', 'updated_at'])

        ActivityLog.objects.create(
            playbook=playbook,
            user=user,
            action="OpenTide metadata refreshed",
            details="Metadata refreshed from workbench fields",
        )

        metadata = playbook.opentide_yaml.get('metadata', {}) if playbook.opentide_yaml else {}
        return RefreshOpenTideMetadata(
            success=True,
            message="Metadata refreshed successfully",
            metadata=json.dumps(metadata),
        )


class Mutation(graphene.ObjectType):
    admin_approve_deployment = AdminApproveDeployment.Field()
    # Playbook & Graph CRUD
    create_playbook = CreatePlaybook.Field()
    update_playbook_links = UpdatePlaybookLinks.Field()
    update_playbook_tags = UpdatePlaybookTags.Field()
    create_task = CreateTask.Field()
    update_task = UpdateTask.Field()
    delete_task = DeleteTask.Field()

    # Playbook status updates
    update_playbook_status = UpdatePlaybookStatus.Field()
    update_playbook_graph_status = UpdatePlaybookGraphStatus.Field()
    update_own_playbook_graph_status = UpdateOwnPlaybookGraphStatus.Field()
    update_playbook_graph_title = UpdatePlaybookGraphTitle.Field()

    # Core metadata updates
    update_playbook = UpdatePlaybook.Field()
    update_playbook_details = UpdatePlaybookDetails.Field()
    create_capability_abstraction = CreateCapabilityAbstraction.Field()
    update_capability_abstraction = UpdateCapabilityAbstraction.Field()
    delete_capability_abstraction = DeleteCapabilityAbstraction.Field()
    submit_for_review = SubmitForReview.Field()
    finalize_review = FinalizeReview.Field()

    # Graph metadata & snapshot
    update_playbook_graph_metadata = UpdatePlaybookGraphMetadata.Field()
    set_playbook_graph_snapshot = SetPlaybookGraphSnapshot.Field()
    upload_graph_snapshot = UploadGraphSnapshot.Field()

    # Deletion
    delete_playbook_graph = DeletePlaybookGraph.Field()
    delete_detection_playbook = DeleteDetectionPlaybook.Field()

    # V2 graph mutations
    create_playbook_graph = CreatePlaybookGraph.Field()
    create_playbook_node = CreatePlaybookNode.Field()
    create_playbook_edge = CreatePlaybookEdge.Field()
    update_playbook_node_position = UpdatePlaybookNodePosition.Field()
    delete_playbook_node = DeletePlaybookNode.Field()
    delete_playbook_edge = DeletePlaybookEdge.Field()

    # Node updates
    update_node_template = UpdateNodeTemplate.Field()
    update_playbook_node_layer_name = UpdatePlaybookNodeLayerName.Field()
    update_node_attack_mappings = UpdateNodeAttackMappings.Field()

    # Sharing & cloning
    share_playbook = SharePlaybook.Field()
    share_playbook_graph = SharePlaybookGraph.Field()
    clone_playbook_graph = ClonePlaybookGraph.Field()
    clone_playbook = ClonePlaybook.Field()

    # Upload graph snapshot
    upload_graph_snapshot = UploadGraphSnapshot.Field()

    # Playbook details (strategy & context)
    update_playbook_details = UpdatePlaybookDetails.Field()
    update_playbook_node_color = UpdatePlaybookNodeColor.Field()
    
    # Comments
    add_playbook_comment = AddPlaybookComment.Field()

    # Review submission
    submit_for_review = SubmitForReview.Field()

    # Finalize review
    finalize_review = FinalizeReview.Field()

    # Export/Import playbook
    export_playbook_graph = ExportPlaybookGraph.Field()
    export_workbench_document = ExportWorkbenchDocument.Field()
    export_all_workbenches_hex_v2 = ExportAllWorkbenchesHexV2.Field()
    import_playbook_graph = ImportPlaybookGraph.Field()
    
    # GitHub integration
    push_playbook_to_github = PushPlaybookToGitHub.Field()
    pull_playbook_from_github = PullPlaybookFromGitHub.Field()

    # OpenTide integration
    update_playbook_opentide_yaml = UpdatePlaybookOpenTideYaml.Field()
    refresh_opentide_metadata = RefreshOpenTideMetadata.Field()
    import_from_opentide = ImportFromOpenTide.Field()
